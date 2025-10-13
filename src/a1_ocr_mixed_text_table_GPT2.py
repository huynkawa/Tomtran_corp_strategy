# -*- coding: utf-8 -*-
"""
a2_ocr_mixed_text_table_GPT.py
---------------------------------------
Đọc & xử lý tài liệu hỗn hợp (DOCX / PDF / IMG / TXT / Excel):
- Tự nhận dạng loại file, OCR khi cần (xử lý RAM, không lưu .png)
- PDF áp dụng được phạm vi TRANG (--start/--end) thực sự
- Giữ cấu trúc block rõ ràng: [DOCX]/[PDF]/[EXCEL] + [PARA]/[TABLE]/[SHEET]
- BẢNG xuất theo TSV (tab '\t') để "vector-ready"
- GPT tham gia theo YAML:
    + table_only: chỉ chuẩn hoá bảng (không bịa số)
    + paragraph_with_headings: chuẩn hoá tiêu đề, KHÔNG diễn giải
- Auto-fix heading nhẹ dựa trên YAML (nếu có heading_patterns)
- Xuất: <name>_text.txt + <name>_meta.json
  (tuỳ chọn: --vector-jsonl để xuất thêm <name>_vector.jsonl)

Author: TOMTRAN (revised, portable & structured)
"""

import os, re, io, json, hashlib, argparse, sys
from pathlib import Path
from typing import List, Dict, Optional, Tuple

import pdfplumber
from PIL import Image, ImageOps
import pytesseract
from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from tqdm import tqdm
from langdetect import detect
import yaml
import pandas as pd

# ====== ENV / GPT enhancer ======
# Lưu ý: module src.env sẽ nạp OPENAI_API_KEY từ .env.active nếu có
try:
    import src.env  # noqa: F401
except Exception:
    pass

# 1 điểm import duy nhất (tránh trùng lặp)
try:
    from src.gpt_enhancer import enhance_with_gpt as _enhance_with_gpt
except Exception:
    def _enhance_with_gpt(text, meta=None, image=None, **kwargs):
        # fallback "an toàn": không dùng GPT
        return text

# ====== Cấu hình mặc định (portable, có thể override bằng CLI) ======
OCR_LANG_DEFAULT = "vie+eng"
OCR_PSM_DEFAULT = "6"  # 6 = Assume a single uniform block of text
OCR_CFG_TEMPLATE = "--psm {psm} preserve_interword_spaces=1"

# Mặc định portable (không hard-code D:\). Có thể override bằng CLI.

INPUT_DIR  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\a_text_only_inputs_test"
OUTPUT_DIR = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\a1_ocr_mixed_text_table_GPT"
YAML_RULE_PATH= r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_text_only_rules.yaml"

# ====== Tiện ích chung ======
def sha1_of_text(text: str) -> str:
    return hashlib.sha1((text or "").encode("utf-8")).hexdigest()

def detect_language_safe(text: str) -> str:
    try:
        return detect(text or "")
    except Exception:
        return "unknown"

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)

def load_yaml_rules(path: str) -> dict:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        return {}

def match_yaml_meta(file_name: str, rules: dict) -> dict:
    """
    Tìm metadata mặc định theo regex từ YAML:
    rules:
      defaults: {...}
      files:
        - match: "regex"
          company: "UIC"
          appendix_id: "App 21"
          ...
    """
    if not rules:
        return {}
    defaults = rules.get("defaults", {}) or {}
    for r in rules.get("files", []) or []:
        patt = r.get("match")
        if patt and re.search(patt, file_name, flags=re.I):
            one = r.copy()
            one.pop("match", None)
            return {**defaults, **one}
    return defaults

def build_gpt_prompt_from_yaml(yaml_rules: dict, mode: str) -> str:
    """
    Kết hợp prompt từ YAML theo mode:
      - table_only
      - paragraph_with_headings
    """
    prompt = []
    gp = (yaml_rules or {}).get("gpt_prompt") or {}
    # prompt chung
    if gp.get("common"):
        prompt.append(str(gp["common"]).strip())
    # prompt theo mode
    if mode and gp.get(mode):
        prompt.append(str(gp[mode]).strip())

    # Ràng buộc an toàn
    policy = (yaml_rules or {}).get("policy") or {}
    if policy.get("no_hallucination", True):
        prompt.append("Tuyệt đối không bịa đặt nội dung/số liệu; không suy diễn ngoài văn bản gốc.")
    if policy.get("keep_units", True):
        prompt.append("Không tự đổi đơn vị; giữ nguyên đơn vị và số liệu như gốc.")
    if policy.get("no_translation", True):
        prompt.append("Không dịch thuật ngữ chuyên ngành; giữ nguyên ngôn ngữ gốc.")

    # Chuẩn TSV cho bảng
    if mode == "table_only":
        prompt.append("Đầu ra BẢNG theo TSV: mỗi ô cách nhau bằng tab (\\t), mỗi hàng một dòng. Không thêm mô tả.")
    # Paragraph
    if mode == "paragraph_with_headings":
        prompt.append("Chỉ CHUẨN HOÁ tiêu đề/heading & dàn ý; không diễn giải thêm nội dung.")

    return "\n".join([p for p in prompt if p]).strip()

def normalize_to_tsv(rows: List[List[str]]) -> str:
    """
    Nhận list 2D và trả về TSV (tab-delimited).
    """
    out_lines = []
    for row in rows:
        safe_cells = [(str(c) if c is not None else "").strip() for c in row]
        out_lines.append("\t".join(safe_cells))
    return "\n".join(out_lines)

def is_tableish_line(line: str) -> bool:
    """
    Heuristic: 1 dòng có nhiều khoảng cách/cột -> coi là dòng bảng.
    """
    # Có nhiều tab hoặc có '|' hoặc có >= 3 nhóm khoảng trắng dài
    return ("\t" in line) or ("|" in line) or (len(re.findall(r"\s{2,}", line)) >= 2)

def split_text_into_text_vs_table_blocks(text: str) -> List[Dict[str, str]]:
    """
    Từ text layer PDF: tách block TABLE vs TEXT thô bằng heuristic (mềm).
    """
    blocks, buf, buf_is_table = [], [], None
    for raw in (text or "").splitlines():
        line = raw.rstrip()
        line_is_table = is_tableish_line(line)
        if buf_is_table is None:
            buf_is_table = line_is_table
            buf = [line]
        elif line_is_table == buf_is_table:
            buf.append(line)
        else:
            blocks.append({"type": "table" if buf_is_table else "paragraph",
                           "content": "\n".join(buf).strip()})
            buf = [line]
            buf_is_table = line_is_table
    if buf:
        blocks.append({"type": "table" if buf_is_table else "paragraph",
                       "content": "\n".join(buf).strip()})
    return blocks

# ====== DOCX ======
def read_docx_paragraphs_and_tables(file_path: Path) -> List[Dict[str, str]]:
    """
    Đọc DOCX với cấp cao:
    - paragraphs: lấy style.name để suy ra Heading
    - tables: lấy tất cả cell theo hàng/cột -> TSV
    """
    doc = Document(file_path)
    results: List[Dict[str, str]] = []

    # 1) Paragraphs
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = getattr(para.style, "name", "") or ""
        # Đưa heading ra trước nội dung (giữ làm mốc)
        if style_name.lower().startswith("heading"):
            results.append({"type": "paragraph", "content": f"{text}"})
        else:
            results.append({"type": "paragraph", "content": text})

    # 2) Tables
    for t in doc.tables:
        rows2d: List[List[str]] = []
        for r in t.rows:
            row_cells = []
            for c in r.cells:
                row_cells.append((c.text or "").strip())
            # một số DOCX lặp lại cell do merge; loại bỏ trùng liên tiếp
            dedup = []
            for i, val in enumerate(row_cells):
                if i == 0 or val != row_cells[i-1]:
                    dedup.append(val)
            rows2d.append(dedup)
        table_tsv = normalize_to_tsv(rows2d)
        results.append({"type": "table", "content": table_tsv})

    return results

# ====== PDF ======
def preprocess_pil_for_ocr(img: Image.Image) -> Image.Image:
    """
    Tiền xử lý đơn giản trước khi OCR: chuyển L + autocontrast.
    (Tránh OpenCV để giảm phụ thuộc.)
    """
    gray = img.convert("L")
    gray = ImageOps.autocontrast(gray)
    return gray

def read_pdf_text_and_images(file_path: Path, page_start: int = 1, page_end: Optional[int] = None,
                             ocr_lang: str = OCR_LANG_DEFAULT, ocr_psm: str = OCR_PSM_DEFAULT) -> List[Dict[str, str]]:
    """
    Đọc PDF: áp dụng phạm vi trang, tách TEXT layer và OCR ảnh.
    - TEXT: tách block paragraph/table nhẹ bằng heuristic
    - IMG : OCR -> coi là table block (để GPT chuẩn hoá bảng)
    """
    results: List[Dict[str, str]] = []
    if page_start < 1:
        page_start = 1
    try:
        with pdfplumber.open(file_path) as pdf:
            N = len(pdf.pages)
            if not page_end or page_end > N:
                page_end = N
            # 1-based inclusive -> zero-based slice
            for idx in range(page_start - 1, page_end):
                page = pdf.pages[idx]
                text = page.extract_text() or ""
                if text.strip():
                    # tách block TABLE/TEXT dựa heuristic
                    for b in split_text_into_text_vs_table_blocks(text):
                        # nhãn nguồn trang
                        btype = b["type"]
                        content = b["content"]
                        if content:
                            results.append({
                                "type": btype,
                                "content": content,
                                "page": idx + 1
                            })
                # OCR ảnh (mỗi ảnh -> table block)
                for im in page.images or []:
                    x0, top, x1, bottom = im["x0"], im["top"], im["x1"], im["bottom"]
                    bbox = (x0, top, x1, bottom)
                    cropped_page = page.crop(bbox)
                    try:
                        pil_img = cropped_page.to_image(resolution=300).original  # PIL
                        pil_img = preprocess_pil_for_ocr(pil_img)
                        cfg = OCR_CFG_TEMPLATE.format(psm=ocr_psm)
                        txt = pytesseract.image_to_string(pil_img, lang=ocr_lang, config=cfg)
                        if (txt or "").strip():
                            results.append({
                                "type": "table",
                                "content": txt.strip(),
                                "page": idx + 1,
                                "source": "image_ocr"
                            })
                    except Exception as e:
                        print(f"⚠️ OCR ảnh lỗi (page {idx+1}): {e}")
    except Exception as e:
        print(f"⚠️ PDF đọc lỗi {file_path}: {e}")
    return results

# ====== IMG ======
def ocr_image_to_text(file_path: Path, ocr_lang: str, ocr_psm: str) -> str:
    img = Image.open(file_path)
    pil = preprocess_pil_for_ocr(img)
    cfg = OCR_CFG_TEMPLATE.format(psm=ocr_psm)
    return pytesseract.image_to_string(pil, lang=ocr_lang, config=cfg)

# ====== TXT ======
def read_txt(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

# ====== Excel ======
def read_excel_as_tsv_blocks(file_path: Path) -> List[Dict[str, str]]:
    """
    Mỗi sheet -> TSV (tab-delimited), giữ header nếu có.
    """
    results = []
    try:
        sheets = pd.read_excel(file_path, sheet_name=None)  # header auto-detect
        for sheet_name, df in sheets.items():
            # Đảm bảo chuỗi
            df = df.astype(object).where(pd.notna(df), "")
            # Xuất TSV
            rows2d: List[List[str]] = [list(map(lambda x: str(x).strip(), df.columns.tolist()))]
            for _, row in df.iterrows():
                rows2d.append([str(c).strip() for c in row.tolist()])
            tsv = normalize_to_tsv(rows2d)
            results.append({"sheet": str(sheet_name), "content": tsv})
    except Exception as e:
        print(f"⚠️ Excel đọc lỗi {file_path}: {e}")
    return results

# ====== Heading auto-fix (nhẹ) ======
def apply_heading_autofix(text: str, yaml_rules: dict) -> str:
    """
    Vá heading nhẹ nếu YAML có 'heading_patterns':
      heading_patterns:
        "SECTION D: REFERRAL RISKS": "(?i)\\breferral risks?\\b"
        "APPENDIX I_CFE TARIFF AND TABLE OF CATEGORY": "(?i)\\bCFE\\s+tariff\\b"
    Ý tưởng:
      - Nếu tìm thấy pattern nội dung nhưng không có heading chuẩn -> chèn heading trước dòng đầu tiên khớp.
    """
    hp = (yaml_rules or {}).get("heading_patterns") or {}
    if not hp:
        return text
    text_norm = text  # sẽ thao tác trực tiếp
    for heading, patt in hp.items():
        try:
            # đã có heading chuẩn?
            if re.search(re.escape(heading), text_norm, flags=re.I):
                continue
            m = re.search(patt, text_norm, flags=re.I | re.M)
            if m:
                # Chèn heading trước dòng nơi pattern xuất hiện
                start = m.start()
                # Tìm đầu dòng
                line_start = text_norm.rfind("\n", 0, start)
                if line_start == -1:
                    # chèn đầu file
                    text_norm = f"{heading}\n{text_norm}"
                else:
                    insert_pos = line_start + 1
                    text_norm = text_norm[:insert_pos] + f"{heading}\n" + text_norm[insert_pos:]
        except re.error:
            # regex lỗi -> bỏ qua mục này
            continue
    return text_norm

# ====== Vector JSONL (tùy chọn) ======
def append_vector_jsonl(vec_path: Path, content: str, metadata: dict):
    with open(vec_path, "a", encoding="utf-8") as f:
        rec = {"content": content, "metadata": metadata}
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")

# ====== Xử lý chính 1 file ======
def process_file(file_path: Path,
                 yaml_rules: dict,
                 out_dir: Path,
                 ocr_lang: str,
                 ocr_psm: str,
                 page_start: int,
                 page_end: Optional[int],
                 vector_jsonl: bool = False) -> None:
    file_name = file_path.stem
    ext = file_path.suffix.lower()

    # Meta gốc
    meta: Dict = {
        "file": file_name,
        "source_path": str(file_path.resolve()),
        "ocr_lang": ocr_lang,
        "ocr_psm": ocr_psm,
    }
    meta.update(match_yaml_meta(file_name, yaml_rules))

    combined_lines: List[str] = []
    has_table = False
    gpt_applied = False
    gpt_reasons: List[str] = []
    vector_path = out_dir / f"{file_name}_vector.jsonl" if vector_jsonl else None

    # Helper gọi GPT với prompt dựa YAML
    def call_gpt(text: str, mode: str, extra_meta: dict) -> str:
        nonlocal gpt_applied, gpt_reasons
        prompt = build_gpt_prompt_from_yaml(yaml_rules, mode)
        gpt_applied = True
        reason = f"{mode}"
        if extra_meta.get("source"):
            reason += f":{extra_meta['source']}"
        if reason not in gpt_reasons:
            gpt_reasons.append(reason)
        # truyền prompt & mode để enhancer biết
        return _enhance_with_gpt(text, {**meta, **extra_meta, "gpt_mode": mode}, None, prompt=prompt, mode=mode)

    try:
        if ext == ".docx":
            blocks = read_docx_paragraphs_and_tables(file_path)
            # DOCX: block theo thứ tự – para/table
            for idx, b in enumerate(blocks, 1):
                if b["type"] == "table":
                    has_table = True
                    # TSV đã chuẩn -> vẫn cho GPT "table_only" để chuẩn hoá/khử rác
                    enhanced = call_gpt(b["content"], mode="table_only", extra_meta={"doc_type": "DOCX", "table_index": idx})
                    combined_lines.append(f"### [DOCX] [TABLE {idx}]")
                    combined_lines.append(enhanced.strip())
                    if vector_path:
                        append_vector_jsonl(vector_path, enhanced.strip(),
                                            {**meta, "content_type": "TABLE", "doc_type": "DOCX", "table_index": idx})
                else:
                    # Paragraph: cho GPT chuẩn hoá heading nếu muốn (tùy YAML)
                    para = b["content"]
                    if (yaml_rules or {}).get("enable_paragraph_gpt", True):
                        enhanced = call_gpt(para, mode="paragraph_with_headings", extra_meta={"doc_type": "DOCX"})
                    else:
                        enhanced = para
                    combined_lines.append(f"### [DOCX] [PARA]")
                    combined_lines.append(enhanced.strip())
                    if vector_path:
                        append_vector_jsonl(vector_path, enhanced.strip(),
                                            {**meta, "content_type": "TEXT", "doc_type": "DOCX"})

        elif ext == ".pdf":
            blocks = read_pdf_text_and_images(file_path, page_start=page_start, page_end=page_end,
                                              ocr_lang=ocr_lang, ocr_psm=ocr_psm)
            # PDF: block có page + type
            page_table_count: Dict[int, int] = {}
            for b in blocks:
                btype = b.get("type")
                page_no = b.get("page", 0)
                if btype == "table":
                    has_table = True
                    page_table_count[page_no] = page_table_count.get(page_no, 0) + 1
                    idx_on_page = page_table_count[page_no]
                    enhanced = call_gpt(b["content"], mode="table_only",
                                        extra_meta={"doc_type": "PDF", "page": page_no, "table_index": idx_on_page, "source": b.get("source", "text_layer")})
                    combined_lines.append(f"### [PDF page {page_no}] [TABLE {idx_on_page}]")
                    combined_lines.append(enhanced.strip())
                    if vector_path:
                        append_vector_jsonl(vector_path, enhanced.strip(),
                                            {**meta, "content_type": "TABLE", "doc_type": "PDF",
                                             "page": page_no, "table_index": idx_on_page})
                else:
                    # Paragraph: có thể chạy GPT normalize heading (tùy)
                    para = b.get("content", "")
                    if (yaml_rules or {}).get("enable_paragraph_gpt", True):
                        enhanced = call_gpt(para, mode="paragraph_with_headings",
                                            extra_meta={"doc_type": "PDF", "page": page_no})
                    else:
                        enhanced = para
                    combined_lines.append(f"### [PDF page {page_no}] [TEXT]")
                    combined_lines.append(enhanced.strip())
                    if vector_path:
                        append_vector_jsonl(vector_path, enhanced.strip(),
                                            {**meta, "content_type": "TEXT", "doc_type": "PDF", "page": page_no})

        elif ext in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"]:
            # IMG -> OCR rồi GPT
            txt = ocr_image_to_text(file_path, ocr_lang=ocr_lang, ocr_psm=ocr_psm).strip()
            if txt:
                has_table = True
                enhanced = call_gpt(txt, mode="table_only", extra_meta={"doc_type": "IMG"})
                combined_lines.append(f"### [IMG] [TABLE 1]")
                combined_lines.append(enhanced.strip())
                if vector_path:
                    append_vector_jsonl(vector_path, enhanced.strip(),
                                        {**meta, "content_type": "TABLE", "doc_type": "IMG"})

        elif ext in [".txt", ".csv"]:
            raw = read_txt(file_path).strip()
            # TXT -> để nguyên, có thể normalize heading nếu bật
            if (yaml_rules or {}).get("enable_paragraph_gpt", False):
                enhanced = _enhance_with_gpt(raw, {**meta, "gpt_mode": "paragraph_with_headings"}, None,
                                             prompt=build_gpt_prompt_from_yaml(yaml_rules, "paragraph_with_headings"),
                                             mode="paragraph_with_headings")
                gpt_applied = True
                gpt_reasons.append("paragraph_with_headings")
            else:
                enhanced = raw
            combined_lines.append(f"### [TXT] [TEXT]")
            combined_lines.append(enhanced.strip())
            if vector_path:
                append_vector_jsonl(vector_path, enhanced.strip(),
                                    {**meta, "content_type": "TEXT", "doc_type": "TXT"})

        elif ext in [".xlsx", ".xls"]:
            excel_blocks = read_excel_as_tsv_blocks(file_path)
            for i, b in enumerate(excel_blocks, 1):
                has_table = True
                # Cho GPT table_only để giữ TSV sạch
                enhanced = call_gpt(b["content"], mode="table_only",
                                    extra_meta={"doc_type": "EXCEL", "sheet": b["sheet"]})
                combined_lines.append(f"### [EXCEL] [SHEET={b['sheet']}] [TABLE {i}]")
                combined_lines.append(enhanced.strip())
                if vector_path:
                    append_vector_jsonl(vector_path, enhanced.strip(),
                                        {**meta, "content_type": "TABLE", "doc_type": "EXCEL", "sheet": b["sheet"], "table_index": i})

        else:
            print(f"⚠️ Không hỗ trợ định dạng {ext} — bỏ qua: {file_path.name}")
            return

        combined_text = "\n".join(combined_lines).strip()

        # Auto-fix heading nhẹ nếu YAML có heading_patterns
        combined_text = apply_heading_autofix(combined_text, yaml_rules)

        # Enrich meta
        lang = detect_language_safe(combined_text)
        meta.update({
            "language": lang,
            "has_table": has_table,
            "gpt_applied": gpt_applied,
            "gpt_reasons": sorted(set(gpt_reasons)),
            "text_sha1": sha1_of_text(combined_text),
            "rules_version": (yaml_rules.get("version") if yaml_rules else None),
            "page_range_applied": {"start": page_start, "end": page_end} if page_start or page_end else None
        })

        # Ghi output
        txt_out = out_dir / f"{file_name}_text.txt"
        meta_out = out_dir / f"{file_name}_meta.json"
        with open(txt_out, "w", encoding="utf-8") as f:
            f.write(combined_text)
        with open(meta_out, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

        print(f"✅ Saved: {txt_out.name}, {meta_out.name}{' + vector.jsonl' if vector_path else ''}")

    except Exception as e:
        print(f"❌ Lỗi xử lý {file_path.name}: {e}")

# ====== CLI ======
def build_argparser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="A2 — OCR + GPT cho tài liệu hỗn hợp (text + bảng + ảnh), xuất vector-ready TSV.")
    p.add_argument("--input", default=INPUT_DIR, help="Thư mục input (mặc định: ./inputs)")
    p.add_argument("--out", default=OUTPUT_DIR, help="Thư mục output (mặc định: ./outputs/a2_ocr_mixed_text_table_GPT)")
    p.add_argument("--yaml", default=YAML_RULE_PATH, help="Đường dẫn YAML rule (mặc định: ./configs/a1_text_only_rules.yaml)")
    p.add_argument("--mode", choices=["pages", "files"], default="pages",
                   help="pages = --start/--end là phạm vi TRANG cho PDF; files = theo chỉ số FILE")
    p.add_argument("--start", type=int, default=1, help="[pages] TRANG bắt đầu (1-based)")
    p.add_argument("--end", type=int, default=None, help="[pages] TRANG kết thúc (1-based, inclusive). Nếu bỏ trống -> đến hết")
    p.add_argument("--file_index", type=int, default=None,
                   help="[pages] Chỉ chạy 1 PDF theo thứ tự (1-based). Bỏ qua để chạy tất cả PDF. Không ảnh hưởng non-PDF.")
    p.add_argument("--ocr-lang", default=OCR_LANG_DEFAULT, help="Ngôn ngữ OCR cho ảnh/PDF image (vd: vie+eng)")
    p.add_argument("--ocr-psm", default=OCR_PSM_DEFAULT, help="Tesseract PSM (vd: 4/6/11...)")
    p.add_argument("--vector-jsonl", action="store_true", help="Xuất thêm <name>_vector.jsonl (mỗi block 1 dòng)")
    return p

def main():
    args = build_argparser().parse_args()

    input_dir = Path(args.input).resolve()
    out_dir   = Path(args.out).resolve()
    ensure_dir(out_dir)

    yaml_rules = load_yaml_rules(args.yaml)
    files = sorted(list(input_dir.rglob("*.*")))

    print("=== CẤU HÌNH A2 (revised) ===")
    print(f"📂 INPUT_DIR   : {input_dir}")
    print(f"📦 OUTPUT_DIR  : {out_dir}")
    print(f"⚙️ YAML_RULE   : {args.yaml}  | loaded={'OK' if yaml_rules else 'empty'}")
    print(f"🧠 MODE        : {args.mode}")
    if args.mode == "pages":
        print(f"📄 PAGES       : {args.start} → {args.end or 'END'}")
        print(f"🔢 file_index  : {args.file_index or 'ALL PDFs'}")
    else:
        print(f"🔢 FILE index  : {args.start} → {args.end or len(files)}")
    print(f"🔤 OCR_LANG    : {args.ocr_lang} | PSM={args.ocr_psm}")
    print(f"🧾 VECTOR.JSONL: {'ON' if args.vector_jsonl else 'OFF'}")
    print("=============================")

    if args.mode == "pages":
        pdfs = [Path(f) for f in files if Path(f).suffix.lower() == ".pdf"]
        if args.file_index and args.file_index > 0 and args.file_index <= len(pdfs):
            target = pdfs[args.file_index - 1]
            pdfs = [target]
        elif args.file_index and args.file_index > len(pdfs):
            print(f"⚠️ file_index={args.file_index} > số PDF ({len(pdfs)}). Sẽ chạy toàn bộ PDF.")

        # 1) PDF theo range trang
        for fpath in tqdm(pdfs, desc="Processing PDFs by pages"):
            process_file(
                file_path=Path(fpath),
                yaml_rules=yaml_rules,
                out_dir=out_dir,
                ocr_lang=args.ocr_lang,
                ocr_psm=args.ocr_psm,
                page_start=args.start,
                page_end=args.end,
                vector_jsonl=args.vector_jsonl
            )
        # 2) Non-PDF full
        others = [Path(f) for f in files if Path(f).suffix.lower() != ".pdf"]
        if others:
            for fpath in tqdm(others, desc="Processing non-PDFs full"):
                process_file(
                    file_path=Path(fpath),
                    yaml_rules=yaml_rules,
                    out_dir=out_dir,
                    ocr_lang=args.ocr_lang,
                    ocr_psm=args.ocr_psm,
                    page_start=1,
                    page_end=None,
                    vector_jsonl=args.vector_jsonl
                )
    else:
        # MODE=files: --start/--end là chỉ số FILE
        start = args.start if args.start and args.start > 0 else 1
        end = args.end if args.end else len(files)
        selected = files[start - 1: end]
        for fpath in tqdm(selected, desc="Processing by file index"):
            process_file(
                file_path=Path(fpath),
                yaml_rules=yaml_rules,
                out_dir=out_dir,
                ocr_lang=args.ocr_lang,
                ocr_psm=args.ocr_psm,
                page_start=1,
                page_end=None,
                vector_jsonl=args.vector_jsonl
            )
    print("🎯 Hoàn tất A2 (revised). Kiểm tra thư mục output.")
    
if __name__ == "__main__":
    main()
