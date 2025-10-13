# -*- coding: utf-8 -*-
"""
a1_ocr_mixed_text_table_GPT2.py — Dual-YAML (TABLE + TEXT), vector-ready

Cách chạy đơn giản (mặc định đã trỏ D:\):
    python -m src.a1_ocr_mixed_text_table_GPT2 --start 1 --end 2

Chức năng chính:
- Đọc DOCX / PDF / IMG / TXT / Excel
- Áp page-range PDF (--start/--end) thật sự
- Xuất block có nhãn [DOCX]/[PDF]/[EXCEL] + [TEXT]/[TABLE]/[SHEET]
- Bảng → TSV (tab '\t'), “vector-ready”
- GPT điều khiển bởi YAML:
    + TABLE: rules_table_universal.yaml (gpt_prompt.table_only, policy, numeric…)
    + TEXT : a1_ocr_mix_txt_table_GPT.yaml (defaults/files meta, heading_patterns, enable_paragraph_gpt…)
- Gộp 2 YAML (TEXT “đè” meta: defaults, files, heading_patterns, enable_paragraph_gpt)
- Auto-fix heading nhẹ theo heading_patterns (nếu có)
- Xuất: <name>_text.txt + <name>_meta.json (+ <name>_vector.jsonl nếu bật)

Tác giả: bạn & mình 😄
"""

import os, re, io, json, hashlib, argparse
from pathlib import Path
from typing import List, Dict, Optional, Tuple

import pdfplumber
from PIL import Image, ImageOps
import pytesseract
from docx import Document
from tqdm import tqdm
from langdetect import detect
import yaml
import pandas as pd

# ===== ENV & GPT enhancer =====
try:
    import src.env  # nạp OPENAI_API_KEY nếu có
except Exception:
    pass

try:
    from src.gpt_enhancer import enhance_with_gpt as _enhance_with_gpt
except Exception:
    def _enhance_with_gpt(text, meta=None, image=None, **kwargs):
        return text  # fallback an toàn

# ===== Cấu hình mặc định (đã cố định theo yêu cầu) =====
OCR_LANG_DEFAULT = "vie+eng"
OCR_PSM_DEFAULT  = "6"  # single block
OCR_CFG_TEMPLATE = "--psm {psm} preserve_interword_spaces=1"

# Đường dẫn D:\ như bạn yêu cầu
INPUT_DIR  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\a_text_only_inputs_test"
OUTPUT_DIR = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\a1_ocr_mixed_text_table_GPT"
# YAML text-only “cũ” bạn yêu cầu vẫn giữ để tương thích (sẽ dùng như YAML TEXT)
YAML_RULE_PATH = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_text_only_rules.yaml"
# YAML TEXT (mới) bạn nói “ở cùng đường dẫn text only”
YAML_TEXT_PATH_DEFAULT  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_ocr_mix_txt_table_GPT.yaml"
# YAML TABLE (universal table)
YAML_TABLE_PATH_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\rules_table_universal.yaml"

# Nếu bạn chỉ muốn dùng 1 YAML text, cứ để YAML_TEXT_PATH_DEFAULT trỏ đúng file bạn dùng thực tế.
# Script sẽ gộp: TABLE + TEXT (+ YAML_RULE_PATH nếu tồn tại) → TEXT đè meta.


# ===== Tiện ích chung =====
def sha1_of_text(text: str) -> str:
    return hashlib.sha1((text or "").encode("utf-8")).hexdigest()

def detect_language_safe(text: str) -> str:
    try:
        return detect(text or "")
    except Exception:
        return "unknown"

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)

def load_yaml_rules(path: Optional[str]) -> dict:
    if not path:
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        return {}

def deep_merge(a: dict, b: dict) -> dict:
    """
    Gộp dict b vào a (đệ quy). List thì nối; dict thì merge; kiểu khác -> lấy b.
    """
    if not isinstance(a, dict) or not isinstance(b, dict):
        return b
    out = dict(a)
    for k, v in b.items():
        if k in out:
            if isinstance(out[k], dict) and isinstance(v, dict):
                out[k] = deep_merge(out[k], v)
            elif isinstance(out[k], list) and isinstance(v, list):
                out[k] = out[k] + v
            else:
                out[k] = v
        else:
            out[k] = v
    return out

def match_yaml_meta(file_name: str, rules: dict) -> dict:
    """
    rules:
      defaults: {...}
      files:
        - match: "regex"
          company: "UIC"
          appendix_id: "Appendix 21"
          class: "uw"
    """
    if not rules:
        return {}
    defaults = rules.get("defaults", {}) or {}
    for r in rules.get("files", []) or []:
        patt = r.get("match")
        if patt and re.search(patt, file_name, flags=re.I):
            one = r.copy()
            one.pop("match", None)
            return {**defaults, **one}
    return defaults

# ===== GPT Prompt từ YAML =====
def build_gpt_prompt_from_yaml(yaml_rules: dict, mode: str) -> str:
    gp = (yaml_rules or {}).get("gpt_prompt") or {}
    policy = (yaml_rules or {}).get("policy") or {}

    parts = []
    if gp.get("common"): parts.append(str(gp["common"]).strip())
    if mode and gp.get(mode): parts.append(str(gp[mode]).strip())

    # ràng buộc an toàn
    if policy.get("no_hallucination", True):
        parts.append("• Không bịa/suy diễn nội dung hay số liệu.")
    if policy.get("keep_units", True):
        parts.append("• Không tự đổi đơn vị; giữ nguyên đơn vị và giá trị.")
    if policy.get("no_translation", True):
        parts.append("• Không dịch thuật ngữ chuyên ngành; giữ nguyên ngôn ngữ gốc.")
    if policy.get("allow_reformat_only", True):
        parts.append("• Chỉ tái định dạng/làm sạch/chuẩn hoá; KHÔNG diễn giải.")
    if policy.get("forbid_summary", True):
        parts.append("• Không tóm tắt/nhận xét/diễn giải thêm.")

    if mode == "table_only":
        parts.append("• ĐẦU RA PHẢI LÀ TSV (tab \\t), mỗi hàng một dòng, không mô tả.")

    if mode == "paragraph_with_headings":
        parts.append("• Chỉ chuẩn hoá/khôi phục heading; KHÔNG diễn giải nội dung.")

    return "\n".join([p for p in parts if p]).strip()

# ===== Heuristic: TABLE trong text layer =====
def is_tableish_line(line: str) -> bool:
    return ("\t" in line) or ("|" in line) or (len(re.findall(r"\s{2,}", line)) >= 2)

def split_text_into_text_vs_table_blocks(text: str) -> List[Dict[str, str]]:
    blocks, buf, buf_is_table = [], [], None
    for raw in (text or "").splitlines():
        line = raw.rstrip()
        line_is_table = is_tableish_line(line)
        if buf_is_table is None:
            buf_is_table = line_is_table
            buf = [line]
        elif line_is_table == buf_is_table:
            buf.append(line)
        else:
            blocks.append({"type": "table" if buf_is_table else "paragraph",
                           "content": "\n".join(buf).strip()})
            buf = [line]; buf_is_table = line_is_table
    if buf:
        blocks.append({"type": "table" if buf_is_table else "paragraph",
                       "content": "\n".join(buf).strip()})
    return blocks

def normalize_to_tsv(rows_2d: List[List[str]]) -> str:
    out = []
    for r in rows_2d:
        out.append("\t".join([(str(c) if c is not None else "").strip() for c in r]))
    return "\n".join(out)

# ===== DOCX =====
def read_docx_paragraphs_and_tables(file_path: Path) -> List[Dict[str, str]]:
    doc = Document(file_path)
    results: List[Dict[str, str]] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            results.append({"type": "paragraph", "content": text})
    for t in doc.tables:
        rows: List[List[str]] = []
        for r in t.rows:
            row_cells = [(c.text or "").strip() for c in r.cells]
            # loại trùng liên tiếp do merge
            dedup = []
            for i, val in enumerate(row_cells):
                if i == 0 or val != row_cells[i-1]:
                    dedup.append(val)
            rows.append(dedup)
        results.append({"type": "table", "content": normalize_to_tsv(rows)})
    return results

# ===== PDF =====
def preprocess_pil_for_ocr(img: Image.Image) -> Image.Image:
    gray = img.convert("L")
    return ImageOps.autocontrast(gray)

def read_pdf_text_and_images(
    file_path: Path,
    page_start: int = 1,
    page_end: Optional[int] = None,
    ocr_lang: str = OCR_LANG_DEFAULT,
    ocr_psm: str = OCR_PSM_DEFAULT
) -> List[Dict[str, str]]:
    results: List[Dict[str, str]] = []
    if page_start < 1: page_start = 1
    try:
        with pdfplumber.open(file_path) as pdf:
            total = len(pdf.pages)
            if not page_end or page_end > total:
                page_end = total
            for idx in range(page_start - 1, page_end):
                page = pdf.pages[idx]
                text = page.extract_text() or ""
                if text.strip():
                    for b in split_text_into_text_vs_table_blocks(text):
                        if b["content"]:
                            results.append({"type": b["type"], "content": b["content"], "page": idx + 1})
                # ảnh → OCR → TABLE
                for im in page.images or []:
                    x0, top, x1, bottom = im["x0"], im["top"], im["x1"], im["bottom"]
                    try:
                        pil = page.crop((x0, top, x1, bottom)).to_image(resolution=300).original
                        pil = preprocess_pil_for_ocr(pil)
                        cfg = f"--psm {ocr_psm} preserve_interword_spaces=1"
                        txt = pytesseract.image_to_string(pil, lang=ocr_lang, config=cfg)
                        if txt.strip():
                            results.append({"type": "table", "content": txt.strip(), "page": idx + 1, "source": "image_ocr"})
                    except Exception as e:
                        print(f"⚠️ OCR ảnh lỗi (page {idx+1}): {e}")
    except Exception as e:
        print(f"⚠️ PDF đọc lỗi {file_path.name}: {e}")
    return results

# ===== IMG / TXT / Excel =====
def ocr_image_to_text(file_path: Path, ocr_lang: str, ocr_psm: str) -> str:
    pil = preprocess_pil_for_ocr(Image.open(file_path))
    cfg = f"--psm {ocr_psm} preserve_interword_spaces=1"
    return pytesseract.image_to_string(pil, lang=ocr_lang, config=cfg)

def read_txt(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_excel_as_tsv_blocks(file_path: Path) -> List[Dict[str, str]]:
    results = []
    try:
        sheets = pd.read_excel(file_path, sheet_name=None)
        for sheet, df in sheets.items():
            df = df.astype(object).where(pd.notna(df), "")
            rows = [list(map(lambda x: str(x).strip(), df.columns.tolist()))]
            for _, r in df.iterrows():
                rows.append([str(c).strip() for c in r.tolist()])
            results.append({"sheet": str(sheet), "content": normalize_to_tsv(rows)})
    except Exception as e:
        print(f"⚠️ Excel đọc lỗi {file_path.name}: {e}")
    return results

# ===== Heading auto-fix =====
def apply_heading_autofix(text: str, yaml_rules: dict) -> str:
    hp = (yaml_rules or {}).get("heading_patterns") or {}
    if not hp:
        return text
    out = text
    for heading, patt in hp.items():
        try:
            if re.search(re.escape(heading), out, flags=re.I):
                continue
            m = re.search(patt, out, flags=re.I | re.M)
            if m:
                line_start = out.rfind("\n", 0, m.start())
                insert_pos = 0 if line_start == -1 else line_start + 1
                out = out[:insert_pos] + f"{heading}\n" + out[insert_pos:]
        except re.error:
            continue
    return out

# ===== Class detection (gợi ý) =====
def detect_table_class(block_text: str, yaml_rules: dict, fallback: Optional[str]=None) -> Optional[str]:
    dc = (yaml_rules or {}).get("doc_classification") or {}
    text = (block_text or "").lower()
    def has_any(arr):
        for k in arr:
            if str(k).lower() in text:
                return True
        return False
    if dc.get("financial_keywords") and has_any(dc["financial_keywords"]): return "financial"
    if dc.get("uw_keywords")        and has_any(dc["uw_keywords"]):        return "uw"
    if dc.get("kpi_keywords")       and has_any(dc["kpi_keywords"]):       return "kpi"
    if dc.get("client_keywords")    and has_any(dc["client_keywords"]):    return "client"
    return fallback

# ===== Post-clean (light) =====
def _yaml_bool(d: dict, path: List[str], default: bool) -> bool:
    cur = d or {}
    for k in path:
        if not isinstance(cur, dict) or (k not in cur):
            return default
        cur = cur[k]
    return bool(cur)

def sanitize_text_block(text: str, yaml_rules: dict) -> Tuple[str, List[str]]:
    """Làm sạch nhẹ TEXT sau GPT: bỏ NBSP, '|' lẻ, chuẩn hoá khoảng trắng."""
    warnings = []
    cfg = (yaml_rules or {}).get("text_cleanup", {}) or {}
    out = text or ""

    # strip characters
    for ch in cfg.get("strip_characters", ["\u00A0", "\t", "\r"]):
        out = out.replace(ch, " ")

    # bỏ dấu '|' đơn lẻ (tránh làm hỏng bảng vì TEXT block thôi)
    out = re.sub(r" ?\| ?", " ", out)

    # normalize spaces
    if cfg.get("normalize_spaces", True):
        out = re.sub(r"[ \t]+", " ", out)

    # drop lines matching remove_patterns
    rem_patts = cfg.get("remove_patterns", [])
    if rem_patts:
        kept = []
        for line in out.splitlines():
            drop = False
            for patt in rem_patts:
                try:
                    if re.search(patt, line):
                        drop = True
                        break
                except re.error:
                    continue
            if not drop:
                kept.append(line)
        out = "\n".join(kept)

    # collapse double newlines
    if cfg.get("collapse_double_newlines", True):
        out = re.sub(r"\n{3,}", "\n\n", out.strip())

    if out.strip() == "":
        warnings.append("text_block_empty_after_sanitize")

    return out.strip(), warnings

def sanitize_tsv_block(tsv: str, yaml_rules: dict) -> Tuple[str, List[str]]:
    """Làm sạch nhẹ TSV: trim cell, drop caption/rows rỗng, giữ nguyên giá trị."""
    warnings = []
    rules = yaml_rules or {}
    clean_rules = (rules.get("table_clean_rules") or {})
    validators = (rules.get("validators") or {})

    lines = [ln for ln in (tsv or "").splitlines()]
    new_lines = []
    min_cols = validators.get("min_columns", 1)
    max_cols = validators.get("max_columns", 1000)

    # compile drop patterns
    patt_list = []
    for patt in clean_rules.get("drop_rows_matching", []) or []:
        try:
            patt_list.append(re.compile(patt))
        except re.error:
            continue

    for raw in lines:
        # split by tab (TSV assumed)
        cells = [c.strip() for c in raw.split("\t")]
        # drop caption lines
        drop = False
        for rp in patt_list:
            if rp.search(raw):
                drop = True
                break
        if drop:
            continue

        # drop all-empty row
        if clean_rules.get("drop_if_all_empty", True):
            if all((c == "" for c in cells)):
                continue

        # trim cells always (already trimmed above)
        if clean_rules.get("trim_cells", True):
            cells = [c.strip() for c in cells]

        # basic col bounds
        if len(cells) < min_cols:
            warnings.append(f"row_dropped_min_cols:{len(cells)}<{min_cols}")
            continue
        if len(cells) > max_cols:
            warnings.append(f"row_trimmed_max_cols:{len(cells)}>{max_cols}")
            cells = cells[:max_cols]

        new_lines.append("\t".join(cells))

    cleaned = "\n".join(new_lines).strip()

    # warn if empty
    if cleaned == "":
        warnings.append("table_block_empty_after_sanitize")

    return cleaned, warnings

# ===== Vector JSONL =====
def append_vector_jsonl(vec_path: Path, content: str, metadata: dict):
    with open(vec_path, "a", encoding="utf-8") as f:
        f.write(json.dumps({"content": content, "metadata": metadata}, ensure_ascii=False) + "\n")

# ===== Xử lý 1 file =====
def process_file(
    file_path: Path,
    yaml_rules: dict,
    out_dir: Path,
    ocr_lang: str,
    ocr_psm: str,
    page_start: int,
    page_end: Optional[int],
    vector_jsonl: bool = False
) -> None:
    file_name = file_path.stem
    ext = file_path.suffix.lower()

    meta: Dict = {
        "file": file_name,
        "source_path": str(file_path.resolve()),
        "ocr_lang": ocr_lang,
        "ocr_psm": ocr_psm,
    }
    meta.update(match_yaml_meta(file_name, yaml_rules))

    combined: List[str] = []
    has_table = False
    gpt_applied = False
    gpt_reasons: List[str] = []
    postclean_warnings: List[str] = []

    rules_version = (yaml_rules.get("meta", {}) or {}).get("version")
    vec_path = out_dir / f"{file_name}_vector.jsonl" if vector_jsonl else None

    def call_gpt(text: str, mode: str, extra_meta: dict) -> str:
        nonlocal gpt_applied, gpt_reasons
        prompt = build_gpt_prompt_from_yaml(yaml_rules, mode)
        gpt_applied = True
        reason = mode
        if extra_meta.get("source"):
            reason += f":{extra_meta['source']}"
        if reason not in gpt_reasons:
            gpt_reasons.append(reason)

        # Nhét prompt vào meta để các bản enhancer cũ vẫn đọc được nếu cần
        payload_meta = {**meta, **extra_meta, "gpt_mode": mode, "gpt_prompt": prompt}

        # 1) Thử gọi theo chữ ký mới (có prompt/mode)
        try:
            return _enhance_with_gpt(text, payload_meta, None, prompt=prompt, mode=mode)
        except TypeError:
            # 2) Fallback chữ ký cũ (không có prompt/mode)
            try:
                return _enhance_with_gpt(text, payload_meta, None)
            except TypeError:
                # 3) Fallback tối giản nhất
                return _enhance_with_gpt(text)

    try:
        if ext == ".docx":
            blocks = read_docx_paragraphs_and_tables(file_path)
            tbl_idx = 0
            for b in blocks:
                if b["type"] == "table":
                    has_table = True
                    tbl_idx += 1
                    cls = meta.get("class") or detect_table_class(b["content"], yaml_rules)
                    enhanced = call_gpt(b["content"], "table_only",
                                        {"doc_type": "DOCX", "table_index": tbl_idx, "class": cls})
                    # --- postclean TSV (light)
                    enhanced, warns = sanitize_tsv_block(enhanced, yaml_rules)
                    postclean_warnings.extend([f"docx_table{tbl_idx}:{w}" for w in warns])

                    combined += [f"### [DOCX] [TABLE {tbl_idx}]", enhanced.strip()]
                    if vec_path:
                        append_vector_jsonl(vec_path, enhanced.strip(),
                                            {**meta, "content_type": "TABLE", "doc_type": "DOCX",
                                             "table_index": tbl_idx, "class": cls, "rules_version": rules_version})
                else:
                    para = b["content"]
                    if (yaml_rules or {}).get("enable_paragraph_gpt", True):
                        enhanced = call_gpt(para, "paragraph_with_headings", {"doc_type": "DOCX"})
                    else:
                        enhanced = para
                    # --- postclean TEXT (light)
                    enhanced, warns = sanitize_text_block(enhanced, yaml_rules)
                    postclean_warnings.extend([f"docx_text:{w}" for w in warns])

                    combined += [f"### [DOCX] [TEXT]", enhanced.strip()]
                    if vec_path:
                        append_vector_jsonl(vec_path, enhanced.strip(),
                                            {**meta, "content_type": "TEXT", "doc_type": "DOCX",
                                             "rules_version": rules_version})

        elif ext == ".pdf":
            blocks = read_pdf_text_and_images(file_path, page_start, page_end, ocr_lang, ocr_psm)
            page_table_count: Dict[int, int] = {}
            for b in blocks:
                btype = b.get("type")
                page_no = b.get("page", 0)
                if btype == "table":
                    has_table = True
                    page_table_count[page_no] = page_table_count.get(page_no, 0) + 1
                    idx_on_page = page_table_count[page_no]
                    cls = meta.get("class") or detect_table_class(b.get("content",""), yaml_rules)
                    enhanced = call_gpt(b["content"], "table_only",
                                        {"doc_type":"PDF","page":page_no,"table_index":idx_on_page,
                                         "source": b.get("source","text_layer"), "class": cls})
                    # --- postclean TSV (light)
                    enhanced, warns = sanitize_tsv_block(enhanced, yaml_rules)
                    postclean_warnings.extend([f"pdf_p{page_no}_t{idx_on_page}:{w}" for w in warns])

                    combined += [f"### [PDF page {page_no}] [TABLE {idx_on_page}]", enhanced.strip()]
                    if vec_path:
                        append_vector_jsonl(vec_path, enhanced.strip(),
                                            {**meta, "content_type":"TABLE","doc_type":"PDF",
                                             "page":page_no,"table_index":idx_on_page,
                                             "class": cls, "rules_version": rules_version})
                else:
                    para = b.get("content","")
                    if (yaml_rules or {}).get("enable_paragraph_gpt", True):
                        enhanced = call_gpt(para, "paragraph_with_headings", {"doc_type":"PDF","page":page_no})
                    else:
                        enhanced = para
                    # --- postclean TEXT (light)
                    enhanced, warns = sanitize_text_block(enhanced, yaml_rules)
                    postclean_warnings.extend([f"pdf_p{page_no}_text:{w}" for w in warns])

                    combined += [f"### [PDF page {page_no}] [TEXT]", enhanced.strip()]
                    if vec_path:
                        append_vector_jsonl(vec_path, enhanced.strip(),
                                            {**meta, "content_type":"TEXT","doc_type":"PDF",
                                             "page":page_no, "rules_version": rules_version})

        elif ext in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"]:
            txt = ocr_image_to_text(file_path, ocr_lang, ocr_psm).strip()
            if txt:
                has_table = True
                cls = meta.get("class") or detect_table_class(txt, yaml_rules)
                enhanced = call_gpt(txt, "table_only", {"doc_type":"IMG","class":cls})
                # --- postclean TSV (light)
                enhanced, warns = sanitize_tsv_block(enhanced, yaml_rules)
                postclean_warnings.extend([f"img_table:{w}" for w in warns])

                combined += [f"### [IMG] [TABLE 1]", enhanced.strip()]
                if vec_path:
                    append_vector_jsonl(vec_path, enhanced.strip(),
                                        {**meta, "content_type":"TABLE","doc_type":"IMG",
                                         "class": cls, "rules_version": rules_version})

        elif ext in [".txt", ".csv"]:
            raw = read_txt(file_path).strip()
            if (yaml_rules or {}).get("enable_paragraph_gpt", False):
                enhanced = call_gpt(raw, "paragraph_with_headings", {"doc_type":"TXT"})
            else:
                enhanced = raw
            # --- postclean TEXT (light)
            enhanced, warns = sanitize_text_block(enhanced, yaml_rules)
            postclean_warnings.extend([f"txt_text:{w}" for w in warns])

            combined += [f"### [TXT] [TEXT]", enhanced.strip()]
            if vec_path:
                append_vector_jsonl(vec_path, enhanced.strip(),
                                    {**meta, "content_type":"TEXT","doc_type":"TXT",
                                     "rules_version": rules_version})

        elif ext in [".xlsx", ".xls"]:
            excel_blocks = read_excel_as_tsv_blocks(file_path)
            for i, b in enumerate(excel_blocks, 1):
                has_table = True
                cls = meta.get("class") or detect_table_class(b["content"], yaml_rules)
                enhanced = call_gpt(b["content"], "table_only",
                                    {"doc_type":"EXCEL","sheet":b["sheet"], "table_index": i, "class": cls})
                # --- postclean TSV (light)
                enhanced, warns = sanitize_tsv_block(enhanced, yaml_rules)
                postclean_warnings.extend([f"excel_{b['sheet']}_t{i}:{w}" for w in warns])

                combined += [f"### [EXCEL] [SHEET={b['sheet']}] [TABLE {i}]", enhanced.strip()]
                if vec_path:
                    append_vector_jsonl(vec_path, enhanced.strip(),
                                        {**meta, "content_type":"TABLE","doc_type":"EXCEL",
                                         "sheet": b["sheet"], "table_index": i,
                                         "class": cls, "rules_version": rules_version})

        else:
            print(f"⚠️ Không hỗ trợ định dạng {ext} — bỏ qua: {file_path.name}")
            return

        combined_text = "\n".join(combined).strip()
        combined_text = apply_heading_autofix(combined_text, yaml_rules)

        lang = detect_language_safe(combined_text)
        meta.update({
            "language": lang,
            "has_table": has_table,
            "gpt_applied": gpt_applied,
            "gpt_reasons": sorted(set(gpt_reasons)),
            "text_sha1": sha1_of_text(combined_text),
            "rules_version": rules_version,
            "page_range_applied": {"start": page_start, "end": page_end} if page_start or page_end else None
        })
        if postclean_warnings:
            meta["postclean_warnings"] = postclean_warnings

        ensure_dir(out_dir)
        txt_out  = out_dir / f"{file_name}_text.txt"
        meta_out = out_dir / f"{file_name}_meta.json"
        with open(txt_out, "w", encoding="utf-8") as f:
            f.write(combined_text)
        with open(meta_out, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
        print(f"✅ Saved: {txt_out.name}, {meta_out.name}{' + vector.jsonl' if vec_path else ''}")

    except Exception as e:
        print(f"❌ Lỗi xử lý {file_path.name}: {e}")

# ===== CLI =====
def build_argparser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="A1/A2 — OCR + GPT (Dual-YAML) TSV/vector-ready")
    # Giữ CLI đơn giản: bạn chỉ cần --start/--end; các path có mặc định D:\
    p.add_argument("--start", type=int, default=1, help="[pages] TRANG bắt đầu (1-based)")
    p.add_argument("--end",   type=int, default=None, help="[pages] TRANG kết thúc (1-based, inclusive); None=đến hết")
    p.add_argument("--file_index", type=int, default=None, help="[pages] Chỉ chạy 1 PDF (1-based); bỏ qua = tất cả")
    p.add_argument("--ocr-lang", default=OCR_LANG_DEFAULT, help="Ngôn ngữ OCR (vd: vie+eng)")
    p.add_argument("--ocr-psm",  default=OCR_PSM_DEFAULT, help="PSM Tesseract (vd: 4/6/11)")
    p.add_argument("--no-vector-jsonl", dest="vector_jsonl", action="store_false", help="Tắt xuất <name>_vector.jsonl (mặc định: bật)")
    p.set_defaults(vector_jsonl=True)

    return p

def main():
    args = build_argparser().parse_args()

    input_dir = Path(INPUT_DIR)
    out_dir   = Path(OUTPUT_DIR)

    # Nạp 2 YAML: TABLE + TEXT (+ YAML_RULE_PATH cũ nếu có)
    yaml_table = load_yaml_rules(YAML_TABLE_PATH_DEFAULT)
    yaml_text1 = load_yaml_rules(YAML_TEXT_PATH_DEFAULT)
    yaml_text2 = load_yaml_rules(YAML_RULE_PATH)  # tương thích file a1_text_only_rules.yaml
    yaml_text  = deep_merge(yaml_text1, yaml_text2) if yaml_text2 else yaml_text1

    # Gộp TABLE + TEXT (TEXT đè defaults/files/heading_patterns/enable_paragraph_gpt)
    yaml_rules = deep_merge(yaml_table, yaml_text)

    # Nếu cả 2 có 'files' → ưu tiên YAML TEXT match trước:
    if (yaml_table.get("files") and yaml_text.get("files")):
        yaml_rules["files"] = yaml_text["files"] + yaml_table["files"]

    files = sorted(list(input_dir.rglob("*.*")))

    print("=== CẤU HÌNH (Dual-YAML) ===")
    print(f"📂 INPUT_DIR    : {input_dir}")
    print(f"📦 OUTPUT_DIR   : {out_dir}")
    print(f"📝 YAML_TEXT    : {YAML_TEXT_PATH_DEFAULT} | +compat: {YAML_RULE_PATH}")
    print(f"📐 YAML_TABLE   : {YAML_TABLE_PATH_DEFAULT}")
    print(f"🔤 OCR_LANG     : {args.ocr_lang} | PSM={args.ocr_psm}")
    print(f"📄 PAGES        : {args.start} → {args.end or 'END'} | file_index={args.file_index or 'ALL PDFs'}")
    print(f"🧾 VECTOR.JSONL : {'ON' if args.vector_jsonl else 'OFF'}")
    print("=============================")

    # MODE mặc định: pages (để nhận --start/--end)
    pdfs = [Path(f) for f in files if Path(f).suffix.lower() == ".pdf"]
    if args.file_index and 1 <= args.file_index <= len(pdfs):
        pdfs = [pdfs[args.file_index - 1]]
    elif args.file_index and args.file_index > len(pdfs):
        print(f"⚠️ file_index={args.file_index} > số PDF ({len(pdfs)}). Sẽ chạy toàn bộ PDF.")

    # 1) PDF theo range trang
    for fpath in tqdm(pdfs, desc="Processing PDFs by pages"):
        process_file(
            file_path=Path(fpath),
            yaml_rules=yaml_rules,
            out_dir=out_dir,
            ocr_lang=args.ocr_lang,
            ocr_psm=args.ocr_psm,
            page_start=args.start,
            page_end=args.end,
            vector_jsonl=args.vector_jsonl
        )

    # 2) Non-PDF full
    others = [Path(f) for f in files if Path(f).suffix.lower() != ".pdf"]
    if others:
        for fpath in tqdm(others, desc="Processing non-PDFs full"):
            process_file(
                file_path=Path(fpath),
                yaml_rules=yaml_rules,
                out_dir=out_dir,
                ocr_lang=args.ocr_lang,
                ocr_psm=args.ocr_psm,
                page_start=1,
                page_end=None,
                vector_jsonl=args.vector_jsonl
            )

    print("🎯 Hoàn tất. Kiểm tra thư mục output.")

if __name__ == "__main__":
    main()
