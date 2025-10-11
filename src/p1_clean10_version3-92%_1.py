# -*- coding: utf-8 -*-
"""
p1a_clean10_ocr_bctc_GPT.py
------------------------------------------------
Pipeline OCR BCTC "lai":
- Mặc định dùng OpenCV + Tesseract cho TABLE, YAML clean + validator + narrator.
- Tự động route sang PaddleOCR (PP-Structure) khi trang "khó" (dựa trên score).
- TEXT page: Tesseract → YAML text clean.
- Cross-check GPT (nhánh QA) có thể bật/tắt (mặc định tắt) — không thay thế YAML/validator.

I/O:
- Duyệt file PDF/ảnh/DOCX → chia block theo từng trang.
- Với PDF/ảnh: có PAGE index. Với DOCX: quy ước PAGE tăng dần theo đơn vị khối (paragraph/table) vì DOCX không có page layout cố định.
- Gộp ra 1 TXT + 1 meta.json cho mỗi file input.
- Có thể bật cờ QA để xuất thêm *_TEXT.txt / *_TABLE.txt theo trang.

Chạy mẫu:
  python -m src.p1a_clean10_ocr_bctc_GPT --start 7 --end 8
"""

from __future__ import annotations
import os, re, io, cv2, json, unicodedata, hashlib, argparse
from typing import List, Tuple, Optional, Dict, Any

import numpy as np
from PIL import Image

# --- OCR / PDF / DOCX deps ---
import pytesseract
from pytesseract import Output
from pdf2image import convert_from_path
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# --- Optional: PaddleOCR PP-Structure ---
HAS_PADDLE = False
try:
    # Chỉ khởi tạo khi thật sự dùng, để tránh xung đột numpy/ABI
    from paddleocr import PPStructure, save_structure_res  # noqa
    HAS_PADDLE = True
except Exception:
    HAS_PADDLE = False

# ====== ĐƯỜNG DẪN MẶC ĐỊNH ======
INPUT_ROOT_DEFAULT  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\c_financial_reports"
OUTPUT_ROOT_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\p1a_clean10_ocr_bctc"
YAML_TABLE_DEFAULT  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\p1a_clean10_ocr_bctc_table.yaml"
YAML_TEXT_DEFAULT   = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\p1a_clean10_ocr_bctc_text.yaml"

# ====== Cấu hình chung ======
TESSERACT_LANG = "vie+eng"
TESSERACT_CFG_TEXT = "--psm 6 -c preserve_interword_spaces=1"
TESSERACT_CFG_TSV  = "--psm 6 -c preserve_interword_spaces=1"
DPI = 500
EXT_IMAGE = [".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"]
PADDLE_LANG = "vi"     # hoặc 'en' tùy tài liệu
QA_PREFIX_TEXT  = "_TEXT"
QA_PREFIX_TABLE = "_TABLE"

# ===== Utils =====
def ensure_dir(p): os.makedirs(p, exist_ok=True)
def _sha1(s: str) -> str: return hashlib.sha1((s or "").encode("utf-8")).hexdigest()

def read_yaml_safe(path: str) -> Dict[str, Any]:
    try:
        import yaml
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception as e:
        print(f"⚠️ Không đọc được YAML: {path} → dùng mặc định. Lý do: {e}")
        return {}

def strip_vn_accents(s: str) -> str:
    s = (s or "").replace("Đ","D").replace("đ","d")
    s = unicodedata.normalize("NFD", s)
    return "".join(ch for ch in s if unicodedata.category(ch) != "Mn")

def clean_whitespace(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

# ====== YAML Cleaners (mẫu khả chuyển) ======
def yaml_text_clean(text: str, rules: Dict[str, Any]) -> str:
    out = text or ""
    # áp các regex thay thế (nếu có)
    for pat, rep in (rules.get("regex_replace") or {}).items():
        try:
            out = re.sub(pat, rep, out, flags=re.MULTILINE)
        except re.error:
            pass
    # chuẩn hóa white space
    if rules.get("normalize_whitespace", True):
        out = "\n".join(clean_whitespace(ln) for ln in out.splitlines())
    return out.strip()

def yaml_table_clean(rows: List[Dict[str,str]], rules: Dict[str, Any]) -> List[Dict[str,str]]:
    # ví dụ: map alias, chuẩn hóa dấu/thousand separators, loại dòng nhiễu...
    def _clean_num(x: str) -> str:
        s = (x or "").upper().replace("O","0").replace("o","0").replace("U","0").replace("Ô","0")
        s = s.replace(",", "").replace(" ", "")
        s = re.sub(r"[^0-9\.-]", "", s)
        s = re.sub(r"\.{2,}", ".", s)
        return s.strip(".")
    name_map = rules.get("name_alias", {})
    drop_if_name_contains = rules.get("drop_if_name_contains", [])
    out=[]
    for r in rows:
        ma    = (r.get("ma","") or "").strip()
        chi   = clean_whitespace(r.get("chi","") or "")
        tm    = (r.get("tm","") or "").strip()
        end   = _clean_num(r.get("end","") or "")
        start = _clean_num(r.get("start","") or "")
        low   = strip_vn_accents(chi).lower()
        for k,v in name_map.items():
            if k in low: chi = v
        if any(substr in low for substr in drop_if_name_contains):
            continue
        out.append({"ma":ma,"chi":chi,"tm":tm,"end":end,"start":start})
    return out

# ====== Validator số (mẫu) ======
def validator_fix_numbers(rows: List[Dict[str,str]], rules: Dict[str, Any]) -> List[Dict[str,str]]:
    # Có thể thêm: quan hệ cha-con, tổng thành phần, ép kiểu, fallback "0" nếu rỗng khi bắt buộc...
    force_zero_if_empty = rules.get("force_zero_if_empty", False)
    for r in rows:
        for k in ("end","start"):
            if force_zero_if_empty and not r.get(k):
                r[k] = "0"
    return rows

# ====== Narrator (mô tả cho từng dòng) ======
def narrator_rows(rows: List[Dict[str,str]]) -> str:
    out=[]
    for r in rows:
        bullet = f"- [{r.get('ma','')}] {r.get('chi','')}"
        end, start = r.get("end",""), r.get("start","")
        if end or start:
            bullet += f" — Cuối năm: {end or '∅'}; Đầu năm: {start or '∅'}"
        if r.get("tm"):
            bullet += f" (TM: {r['tm']})"
        out.append(bullet)
    return "\n".join(out)

# ====== OpenCV-based table detection & parsing (rút gọn, ổn định) ======
NUM_TOKEN = re.compile(r"[0-9][0-9\.,]{2,}")
_NUM_GROUP = re.compile(r"\d{1,3}(?:\.\d{3})+")

def preprocess_image_gray_bin(bgr):
    gray  = cv2.cvtColor(bgr, cv2.COLOR_BGR2GRAY)
    blur  = cv2.GaussianBlur(gray, (3,3), 0)
    th    = cv2.adaptiveThreshold(blur,255,cv2.ADAPTIVE_THRESH_GAUSSIAN_C,cv2.THRESH_BINARY,21,8)
    return th

def find_table_bbox(gray: np.ndarray) -> Optional[Tuple[int,int,int,int]]:
    H,W = gray.shape[:2]
    inv = 255 - gray
    horiz_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(20, W//40), 1))
    vert_kernel  = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(20, H//40)))
    horiz = cv2.morphologyEx(inv, cv2.MORPH_OPEN, horiz_kernel, iterations=2)
    vert  = cv2.morphologyEx(inv, cv2.MORPH_OPEN, vert_kernel,  iterations=2)
    table_mask = cv2.add(horiz, vert)
    table_mask = cv2.dilate(table_mask, cv2.getStructuringElement(cv2.MORPH_RECT,(5,5)), iterations=2)
    cnts,_ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not cnts: return None
    areas = [(cv2.contourArea(c), c) for c in cnts]
    areas.sort(key=lambda x: x[0], reverse=True)
    for area, c in areas:
        x,y,w,h = cv2.boundingRect(c)
        if w*h > 0.10*W*H and w > 0.40*W and h > 0.25*H:
            return x,y,w,h
    return None

def is_note_token(text: str) -> bool:
    t = (text or "").strip()
    return bool(re.fullmatch(r"\d{1,2}(?:\.\d)?", t))

def tokens_to_number(tokens: List[Dict]) -> str:
    s = " ".join(t["text"] for t in tokens)
    cands = _NUM_GROUP.findall(s)
    if cands:
        # chọn dạng có nhiều dấu '.' và dài hơn
        def _score(txt): return (txt.count("."), len(re.sub(r"\D","",txt)))
        best = max(cands, key=_score)
        return clean_num(best)
    return clean_num(s)

def clean_num(s: str) -> str:
    s = (s or "").upper().replace("O","0").replace("o","0").replace("U","0").replace("Ô","0")
    s = s.replace(",", "").replace(" ", "")
    s = re.sub(r"[^0-9\.-]", "", s)
    s = re.sub(r"\.{2,}", ".", s)
    return s.strip(".")

_re_code = re.compile(r"^\s*[\+\|\.\-: ]{0,8}(\d{3})\b")
_re_note = re.compile(r"(?:^|\s)(\d{1,2}(?:\.\d)?)\s*$")

def parse_left_side(text: str) -> Tuple[str, str, str]:
    s = (text or "").strip()
    s = re.sub(r"^[\+\|\.\-: ]+", "", s)
    ma = chi = tm = ""
    m = _re_code.match(s)
    if m:
        ma = m.group(1); s = s[m.end():].strip(" .-:|+")
    m = _re_note.search(s)
    if m:
        tm = m.group(1); s = s[:m.start()].strip(" .-:|+")
    chi = s
    return ma, chi, tm

def find_columns_fallback_by_numbers(words: List[Dict], W: int) -> Tuple[int,int]:
    xs=[]
    for w in words:
        if NUM_TOKEN.search(w['text']):
            x = w['left'] + w['width']/2
            if x > W*0.45: xs.append(x)
    if len(xs)>=6:
        xs = np.array(xs, dtype=float)
        c1, c2 = np.percentile(xs,35), np.percentile(xs,70)
        split2 = int((c1+c2)/2.0); split1 = int(min(c1,c2)-50)
    else:
        split1 = int(W*0.58); split2 = int(W*0.76)
    split1 = max(180, split1); split2 = max(split1+120, split2)
    return split1, split2

def reflow_records_table(img_bgr: np.ndarray) -> List[Dict]:
    data = pytesseract.image_to_data(
        img_bgr, lang=TESSERACT_LANG, config=TESSERACT_CFG_TSV, output_type=Output.DICT
    )
    n = len(data['text'])
    words=[]
    for i in range(n):
        t = (data['text'][i] or "").strip()
        if not t: continue
        try: conf = float(data['conf'][i])
        except: conf = 0
        if conf < 0: continue
        words.append({
            "text": t,
            "left": int(data['left'][i]),
            "top": int(data['top'][i]),
            "width": int(data['width'][i]),
            "line_num": int(data.get('line_num',[0]*n)[i]),
            "par_num":  int(data.get('par_num',[0]*n)[i]),
            "block_num":int(data.get('block_num',[0]*n)[i]),
        })
    if not words: return []

    H, W = img_bgr.shape[:2]
    split1, split2 = find_columns_fallback_by_numbers(words, W)

    groups: Dict[Tuple[int,int,int], List[Dict]] = {}
    for w in words:
        key=(w['block_num'], w['par_num'], w['line_num'])
        groups.setdefault(key, []).append(w)

    raw=[]
    for toks in groups.values():
        toks.sort(key=lambda x:(x['top'], x['left']))
        raw.append(toks)
    raw.sort(key=lambda arr: arr[0]['top'])

    MARGIN  = 18
    TM_BAND = 120
    rows=[]
    for toks in raw:
        toks.sort(key=lambda t:t['left'])
        left_tokens=[]; end_tokens=[]; start_tokens=[]
        tm_candidates=[]

        for t in toks:
            xmid = t['left'] + t['width']/2
            if   xmid < (split1 - MARGIN): left_tokens.append(t)
            elif xmid < (split2 - MARGIN): end_tokens.append(t)
            elif xmid > (split2 + MARGIN): start_tokens.append(t)
            else:
                # vùng giữa: đẩy về cột gần nhất
                if abs(xmid - split1) < abs(xmid - split2): end_tokens.append(t)
                else:                                       start_tokens.append(t)
            if (split1 - TM_BAND) <= xmid <= (split1 + MARGIN) and is_note_token(t['text']):
                tm_candidates.append(t['text'])

        left_text = " ".join(t['text'] for t in left_tokens).strip()
        end_tokens_f   = [t for t in end_tokens   if not is_note_token(t['text'])]
        start_tokens_f = [t for t in start_tokens if not is_note_token(t['text'])]

        end_num   = tokens_to_number(end_tokens_f)
        start_num = tokens_to_number(start_tokens_f)

        tm_text = tm_candidates[0] if tm_candidates else ""
        if not tm_text:
            for col_tokens in (end_tokens, start_tokens):
                small_notes = [t for t in col_tokens if is_note_token(t['text'])]
                if small_notes:
                    tm_text = small_notes[0]['text']; break

        ma, chi, tm = parse_left_side(left_text)
        if not tm and tm_text: tm = tm_text
        if ma or chi or tm or end_num or start_num:
            rows.append({"ma": ma, "chi": chi, "tm": tm, "end": end_num, "start": start_num})

    # gộp dòng ngắn tiếp diễn
    fused=[]
    for r in rows:
        if (fused and not r["ma"] and not r["tm"] and not r["end"] and not r["start"]
                and r["chi"] and len(r["chi"]) <= 28):
            fused[-1]["chi"] = (fused[-1]["chi"] + " " + r["chi"]).strip()
        else:
            fused.append(r)
    return fused

def format_ascii_table(rows: List[Dict]) -> str:
    headers = ["Mã số", "Chỉ tiêu", "Thuyết minh", "Số cuối năm", "Số đầu năm"]
    w = {
        "ma": max([len(r.get("ma","")) for r in rows] + [len(headers[0])]) if rows else len(headers[0]),
        "chi":max([len(r.get("chi","")) for r in rows] + [len(headers[1])]) if rows else len(headers[1]),
        "tm": max([len(r.get("tm","")) for r in rows] + [len(headers[2])]) if rows else len(headers[2]),
        "end":max([len(r.get("end","")) for r in rows] + [len(headers[3])]) if rows else len(headers[3]),
        "start":max([len(r.get("start","")) for r in rows] + [len(headers[4])]) if rows else len(headers[4]),
    }
    pad_l = lambda s, ww: (s or "").ljust(ww)
    pad_r = lambda s, ww: (s or "").rjust(ww)
    line = f"+-{'-'*w['ma']}-+-{'-'*w['chi']}-+-{'-'*w['tm']}-+-{'-'*w['end']}-+-{'-'*w['start']}-+"
    out = [line,
           "| "+pad_l(headers[0],w['ma'])+" | "+pad_l(headers[1],w['chi'])+" | "+pad_l(headers[2],w['tm'])+
           " | "+pad_r(headers[3],w['end'])+" | "+pad_r(headers[4],w['start'])+" |",
           line]
    for r in rows:
        out.append("| "+pad_l(r['ma'],w['ma'])+" | "+pad_l(r['chi'],w['chi'])+" | "+pad_l(r['tm'],w['tm'])+
                   " | "+pad_r(r['end'],w['end'])+" | "+pad_r(r['start'],w['start'])+" |")
    out.append(line)
    return "\n".join(out)

# ====== Chấm điểm trang TABLE để route Paddle fallback ======
def score_table_quality(rows: List[Dict]) -> Dict[str, float]:
    if not rows:
        return {"row_coverage":0.0, "missing_amount_ratio":1.0, "dup_amount_ratio":0.0, "score":0.0}

    def is_amount(x): return bool(re.fullmatch(r"-?\d+(?:\.\d+)?", x or ""))

    covered = sum(1 for r in rows if r.get("ma") or r.get("chi"))
    row_coverage = covered / max(1,len(rows))

    miss = sum(1 for r in rows if not (is_amount(r.get("end","")) and is_amount(r.get("start",""))))
    missing_amount_ratio = miss / max(1,len(rows))

    dup = sum(1 for r in rows if r.get("end") and r.get("start") and r["end"] == r["start"])
    dup_amount_ratio = dup / max(1,len(rows))

    # score đơn giản: càng cao càng tốt
    score = row_coverage - 0.5*missing_amount_ratio - 0.3*dup_amount_ratio
    return {
        "row_coverage": row_coverage,
        "missing_amount_ratio": missing_amount_ratio,
        "dup_amount_ratio": dup_amount_ratio,
        "score": score
    }

def need_paddle_fallback(metrics: Dict[str,float]) -> bool:
    violations = 0
    if metrics["row_coverage"] < 0.75: violations += 1
    if metrics["missing_amount_ratio"] > 0.20: violations += 1
    if metrics["dup_amount_ratio"] > 0.15: violations += 1
    return violations >= 2

# ====== Paddle route (an toàn nếu thiếu Paddle) ======
def paddle_extract_rows(img_bgr: np.ndarray, paddle_timeout: int = 25) -> List[Dict]:
    if not HAS_PADDLE:
        # không có Paddle → fallback lại OpenCV TSV
        return reflow_records_table(img_bgr)

    try:
        table_engine = PPStructure(layout=True, show_log=False, lang=PADDLE_LANG)
        # Paddle nhận BGR np.array
        result = table_engine(img_bgr)
        # Lọc các block type=table, nếu có bbox thì cắt và OCR lại bằng TSV của ta để map 5 cột
        table_like = [r for r in result if r.get('type')=='table']
        rows_all=[]
        if not table_like:
            # không nhận ra table → dùng TSV thường
            return reflow_records_table(img_bgr)
        for tb in table_like:
            box = tb.get('bbox') or tb.get('res') and tb['res'].get('bbox')
            if box and len(box)==4:
                x1,y1,x2,y2 = map(int, box)
                x1,y1 = max(0,x1), max(0,y1)
                x2,y2 = min(img_bgr.shape[1],x2), min(img_bgr.shape[0],y2)
                roi = img_bgr[y1:y2, x1:x2].copy()
                rows = reflow_records_table(roi)
                rows_all.extend(rows)
        return rows_all or reflow_records_table(img_bgr)
    except Exception as e:
        print(f"⚠️ Paddle lỗi → dùng TSV: {e}")
        return reflow_records_table(img_bgr)

# ====== TEXT OCR & clean ======
def ocr_text_block(bgr: np.ndarray) -> str:
    txt = pytesseract.image_to_string(bgr, lang=TESSERACT_LANG, config=TESSERACT_CFG_TEXT)
    return txt or ""

# ====== Xác định TEXT vs TABLE page ======
def is_table_like_page(bgr: np.ndarray) -> Tuple[bool, Optional[Tuple[int,int,int,int]]]:
    gray_bin = preprocess_image_gray_bin(bgr)
    bbox = find_table_bbox(gray_bin)
    return (bbox is not None), bbox

# ====== Cross-check GPT (stub an toàn, chỉ log) ======
def gpt_cross_check_image_vs_table(_pil_image: Image.Image, _rows: List[Dict], _openai_key: Optional[str]) -> Dict[str, Any]:
    # Ở đây chỉ stub: hệ thống thực tế sẽ gọi GPT-4o-mini/vision để so ảnh vs bảng,
    # nhưng để tránh chi phí/khóa, ta chỉ trả thống kê cơ bản.
    stats = score_table_quality(_rows)
    return {"status":"skipped_stub", "basic_stats": stats}

# ====== Ghi block theo trang ======
def block_marker(page_no: int, kind: str) -> str:
    return f"### [PAGE {page_no:02d}] [{kind}]"

def save_qa_dump(out_dir: str, base: str, page_no: int, kind: str, content: str):
    ensure_dir(out_dir)
    fname = f"{base}_{kind}_{page_no:02d}.txt"
    with open(os.path.join(out_dir, fname), "w", encoding="utf-8") as f:
        f.write(content or "")

# ====== Xử lý 1 trang ảnh ======
def process_page_image(bgr: np.ndarray,
                       yaml_text_rules: Dict,
                       yaml_table_rules: Dict,
                       table_engine_mode: str = "auto",
                       gpt_cross: bool = False,
                       paddle_timeout: int = 25) -> Tuple[str, Dict]:
    """
    Trả về: (text_block, meta_page)
    - text_block: string đã gồm marker + nội dung
    - meta_page: dict ghi loại, score, route, vv.
    """
    meta = {"type":None, "route":"opencv", "metrics":{}, "paddle_used":False}

    is_table, bbox = is_table_like_page(bgr)
    if not is_table:
        # TEXT page
        meta["type"]="TEXT"
        raw = ocr_text_block(bgr)
        cleaned = yaml_text_clean(raw, yaml_text_rules)
        return cleaned, meta

    # TABLE page
    meta["type"]="TABLE"
    roi = bgr.copy()
    if bbox:
        x,y,w,h = bbox
        roi = bgr[y:y+h, x:x+w]

    # bước 1: TSV (OpenCV+Tesseract)
    rows = reflow_records_table(roi)
    metrics = score_table_quality(rows)
    meta["metrics"] = metrics

    # route auto -> paddle nếu cần
    use_paddle = (table_engine_mode == "paddle") or (table_engine_mode == "auto" and need_paddle_fallback(metrics))
    if use_paddle:
        rows = paddle_extract_rows(roi, paddle_timeout=paddle_timeout)
        meta["route"] = "paddle" if HAS_PADDLE else "opencv"
        meta["paddle_used"] = HAS_PADDLE
        # chấm điểm lại sau paddle
        metrics = score_table_quality(rows)
        meta["metrics"] = metrics
    else:
        meta["route"] = "opencv"

    # YAML clean + validator + narrator
    rows = yaml_table_clean(rows, yaml_table_rules)
    rows = validator_fix_numbers(rows, yaml_table_rules)
    ascii_tbl = format_ascii_table(rows)
    narr = narrator_rows(rows)

    # GPT cross-check (stub)
    if gpt_cross:
        pil_img = Image.fromarray(cv2.cvtColor(roi, cv2.COLOR_BGR2RGB))
        _ = gpt_cross_check_image_vs_table(pil_img, rows, os.environ.get("OPENAI_API_KEY"))

    # TABLE block text = bảng + dòng diễn giải
    block_text = ascii_tbl
    if narr.strip():
        block_text += "\n\n" + narr
    return block_text, meta

# ====== Xử lý theo loại file ======
def process_pdf(path: str, out_dir: str, start: int, end: Optional[int],
                yaml_text_rules: Dict, yaml_table_rules: Dict,
                table_engine_mode: str, qa_dump: bool, gpt_cross: bool, paddle_timeout: int):
    pages = convert_from_path(path, dpi=DPI, first_page=start, last_page=end)
    blocks=[]; metas=[]
    base = os.path.splitext(os.path.basename(path))[0]
    for i, pil in enumerate(pages, start=start):
        bgr = cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)
        content, meta = process_page_image(bgr, yaml_text_rules, yaml_table_rules,
                                           table_engine_mode, gpt_cross, paddle_timeout)
        meta["page"]=i
        kind = meta["type"] or "TEXT"
        blocks.append(block_marker(i, kind) + "\n" + (content or ""))
        metas.append(meta)
        if qa_dump:
            if kind=="TEXT":
                save_qa_dump(out_dir, base+QA_PREFIX_TEXT, i, "TEXT", content or "")
            else:
                save_qa_dump(out_dir, base+QA_PREFIX_TABLE, i, "TABLE", content or "")
    return blocks, metas

def process_image_file(path: str, out_dir: str,
                       yaml_text_rules: Dict, yaml_table_rules: Dict,
                       table_engine_mode: str, qa_dump: bool, gpt_cross: bool, paddle_timeout: int):
    base = os.path.splitext(os.path.basename(path))[0]
    bgr = cv2.imread(path)
    if bgr is None:
        print(f"⚠️ Không đọc được ảnh: {path}")
        return [], []
    content, meta = process_page_image(bgr, yaml_text_rules, yaml_table_rules,
                                       table_engine_mode, gpt_cross, paddle_timeout)
    page_no = 1
    kind = meta["type"] or "TEXT"
    block = block_marker(page_no, kind) + "\n" + (content or "")
    if qa_dump:
        if kind=="TEXT":
            save_qa_dump(out_dir, base+QA_PREFIX_TEXT, page_no, "TEXT", content or "")
        else:
            save_qa_dump(out_dir, base+QA_PREFIX_TABLE, page_no, "TABLE", content or "")
    meta["page"]=page_no
    return [block], [meta]

def process_docx(path: str, out_dir: str,
                 yaml_text_rules: Dict, yaml_table_rules: Dict,
                 table_engine_mode: str, qa_dump: bool, gpt_cross: bool, paddle_timeout: int):
    # DOCX không có "trang" thật sự; ta quy ước PAGE tăng dần theo khối
    if not HAS_DOCX:
        print("⚠️ Chưa cài python-docx → bỏ qua DOCX.")
        return [], []
    doc = Document(path)
    blocks=[]; metas=[]
    base = os.path.splitext(os.path.basename(path))[0]
    page_no = 0

    # 1) Tables trong DOCX → rasterize từng bảng thành ảnh (via PIL từ screenshot tạm thời)
    # Đơn giản hoá: đọc text từng cell, nối thành ảnh trắng đen bằng font mặc định → OCR lại
    # Nhưng để ngắn gọn và ổn định, ở đây: ta coi bảng DOCX như TEXT khối "TABLE-TEXT"
    # (Nếu cần chính xác cao, nên render DOCX sang PDF trước rồi chạy cùng pipeline PDF.)
    for t in doc.tables:
        page_no += 1
        # tạo văn bản thô từ bảng
        lines=[]
        for row in t.rows:
            cells = [clean_whitespace(c.text) for c in row.cells]
            lines.append(" | ".join(cells))
        raw = "\n".join(lines)
        # coi là TABLE -> đưa qua YAML table sau khi map sơ bộ
        # map rất đơn giản: không có tọa độ → cố gắng parse CODE ở đầu cell[0]
        rows=[]
        for ln in lines:
            left, *nums = [seg.strip() for seg in ln.split("|")]
            ma, chi, tm = parse_left_side(left)
            end = nums[0] if len(nums)>0 else ""
            start = nums[1] if len(nums)>1 else ""
            rows.append({"ma":ma,"chi":chi,"tm":tm,"end":clean_num(end),"start":clean_num(start)})
        rows = yaml_table_clean(rows, yaml_table_rules)
        rows = validator_fix_numbers(rows, yaml_table_rules)
        ascii_tbl = format_ascii_table(rows)
        narr = narrator_rows(rows)
        content = ascii_tbl + ("\n\n"+narr if narr else "")
        kind = "TABLE"
        blocks.append(block_marker(page_no, kind)+"\n"+content)
        metas.append({"type":kind,"page":page_no,"route":"docx-raw","metrics":score_table_quality(rows)})
        if qa_dump:
            save_qa_dump(out_dir, base+QA_PREFIX_TABLE, page_no, "TABLE", content or "")

    # 2) Paragraph TEXT
    buf=[]
    for p in doc.paragraphs:
        s = clean_whitespace(p.text)
        if s: buf.append(s)
    if buf:
        page_no += 1
        raw_text = "\n".join(buf)
        txt = yaml_text_clean(raw_text, yaml_text_rules)
        kind="TEXT"
        blocks.append(block_marker(page_no, kind)+"\n"+txt)
        metas.append({"type":kind,"page":page_no,"route":"docx-text"})
        if qa_dump:
            save_qa_dump(out_dir, base+QA_PREFIX_TEXT, page_no, "TEXT", txt or "")

    return blocks, metas

# ====== Gộp & lưu kết quả cho 1 input file ======
def save_merged_outputs(blocks: List[str], metas: List[Dict], out_dir: str, base: str):
    ensure_dir(out_dir)
    text = "\n\n".join(blocks).strip()
    with open(os.path.join(out_dir, f"{base}.txt"), "w", encoding="utf-8") as f:
        f.write(text)
    meta = {
        "text_sha1": _sha1(text),
        "num_blocks": len(blocks),
        "pages": metas
    }
    with open(os.path.join(out_dir, f"{base}.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

# ====== Batch driver ======
def batch_process(input_root: str,
                  output_root: str,
                  start: int,
                  end: Optional[int],
                  yaml_text_path: str,
                  yaml_table_path: str,
                  table_engine_mode: str,
                  qa_dump: bool,
                  gpt_cross: bool,
                  paddle_timeout: int,
                  paddle_max_pages: int):
    yaml_text_rules  = read_yaml_safe(yaml_text_path)
    yaml_table_rules = read_yaml_safe(yaml_table_path)

    for root, _, files in os.walk(input_root):
        files = sorted(files)
        for fname in files:
            path = os.path.join(root, fname)
            ext  = os.path.splitext(fname)[1].lower()
            base = os.path.splitext(fname)[0]
            print(f"📄 Input: {path}")
            blocks=[]; metas=[]
            # xử lý từng loại
            if ext == ".pdf":
                b, m = process_pdf(path, output_root, start, end,
                                   yaml_text_rules, yaml_table_rules,
                                   table_engine_mode, qa_dump, gpt_cross, paddle_timeout)
                blocks.extend(b); metas.extend(m)
            elif ext in EXT_IMAGE:
                b, m = process_image_file(path, output_root,
                                          yaml_text_rules, yaml_table_rules,
                                          table_engine_mode, qa_dump, gpt_cross, paddle_timeout)
                blocks.extend(b); metas.extend(m)
            elif ext == ".docx":
                b, m = process_docx(path, output_root,
                                    yaml_text_rules, yaml_table_rules,
                                    table_engine_mode, qa_dump, gpt_cross, paddle_timeout)
                blocks.extend(b); metas.extend(m)
            else:
                print(f"ℹ️ Bỏ qua định dạng không hỗ trợ: {fname}")
                continue

            if not blocks:
                print("⚠️ Không có block nào được tạo — bỏ qua lưu.")
                continue
            save_merged_outputs(blocks, metas, output_root, base)
            print(f"✅ Đã lưu: {os.path.join(output_root, base + '.txt')}")

# ====== Main / CLI ======
def main():
    p = argparse.ArgumentParser("BCTC OCR Pipeline (hybrid engine)")
    p.add_argument("--input",  default=INPUT_ROOT_DEFAULT,  help="Thư mục chứa PDF/ảnh/DOCX")
    p.add_argument("--output", default=OUTPUT_ROOT_DEFAULT, help="Thư mục lưu kết quả")
    p.add_argument("--yaml-table", default=YAML_TABLE_DEFAULT, help="YAML rules cho TABLE")
    p.add_argument("--yaml-text",  default=YAML_TEXT_DEFAULT,  help="YAML rules cho TEXT")
    p.add_argument("--start", type=int, default=1, help="Trang bắt đầu của PDF")
    p.add_argument("--end",   type=int, default=None, help="Trang kết thúc của PDF")
    p.add_argument("--table-engine", choices=["auto","opencv","paddle"], default="auto",
                   help="Chọn engine TABLE: opencv (mặc định), paddle, hoặc auto (route theo score)")
    p.add_argument("--qa-dump", choices=["y","n"], default="n", help="Xuất *_TEXT/_TABLE dump theo trang")
    p.add_argument("--gpt-cross-check", choices=["y","n"], default="n",
                   help="Bật cross-check GPT (stub logging, không gọi OpenAI thực)")
    p.add_argument("--paddle-timeout", type=int, default=25, help="Timeout Paddle mỗi trang (giây)")
    p.add_argument("--paddle-max-pages", type=int, default=50, help="Giới hạn số trang route sang Paddle")
    a = p.parse_args()

    if a.table_engine == "paddle" and not HAS_PADDLE:
        print("⚠️ Bạn chọn --table-engine paddle nhưng PaddleOCR chưa sẵn sàng → sẽ dùng opencv.")

    batch_process(
        input_root=a.input,
        output_root=a.output,
        start=a.start,
        end=a.end,
        yaml_text_path=a.yaml_text,
        yaml_table_path=a.yaml_table,
        table_engine_mode=a.table_engine,
        qa_dump=(a.qa_dump=="y"),
        gpt_cross=(a.gpt_cross_check=="y"),
        paddle_timeout=a.paddle_timeout,
        paddle_max_pages=a.paddle_max_pages
    )
    print("🎉 Hoàn tất!")

if __name__ == "__main__":
    main()
