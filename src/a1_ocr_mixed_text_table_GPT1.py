# -*- coding: utf-8 -*-
"""
a2_ocr_mixed_text_table_GPT.py
---------------------------------------
Đọc và xử lý tài liệu hỗn hợp (văn bản + bảng/ảnh):
- Tự nhận dạng loại file: DOCX, PDF, TXT, hình scan.
- OCR khi cần (không lưu .png, xử lý trong RAM).
- GPT chỉ chỉnh sửa phần bảng hoặc ảnh.
- Xuất file text + meta JSON (metadata mở rộng cho vector store).

Giữ nguyên tính năng cũ + THÊM fallback PDF scan (rasterize toàn trang).
"""

import os, re, io, json, hashlib, argparse
from pathlib import Path
from typing import List, Dict, Optional

import fitz  # PyMuPDF
import pdfplumber
from PIL import Image
import pytesseract
from docx import Document
from tqdm import tqdm
from langdetect import detect
import pandas as pd
import yaml

# ===== GPT enhancer (an toàn có fallback) =====
try:
    # nếu chạy kiểu: python -m src.a1_ocr_mixed_text_table_GPT
    from src.gpt_enhancer import enhance_with_gpt as _enhance_with_gpt
except Exception:
    try:
        # nếu file đang trong cùng package
        from .gpt_enhancer import enhance_with_gpt as _enhance_with_gpt
    except Exception:
        # fallback: không dùng GPT, trả nguyên văn để không bị NameError
        def _enhance_with_gpt(text, meta=None, image=None, **kwargs):
            return text

# ========== Cấu hình mặc định (linh hoạt & portable) ==========
BASE_DIR = Path(__file__).resolve().parent.parent

SAVE_PNG_DEBUG = False
OCR_LANG = "vie+eng"
OCR_CFG = "--psm 6 preserve_interword_spaces=1"

# Đường dẫn tuyệt đối (đảm bảo nhất cho pipeline cục bộ)
INPUT_DIR_ABS  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\c_financial_reports_test"
OUTPUT_DIR_ABS = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\a1_ocr_mixed_text_table_GPT_test so sanh p1a"
# Sửa lỗi thiếu dấu cách ở '3. ChatBot_project' nếu người dùng lỡ cấu hình sai
YAML_RULE_PATH_ABS = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_text_only_rules.yaml"

# Fallback nếu đường dẫn tuyệt đối không tồn tại
INPUT_DIR  = INPUT_DIR_ABS if os.path.exists(INPUT_DIR_ABS) else str(BASE_DIR / "inputs" / "c_financial_reports_test")
OUTPUT_DIR = OUTPUT_DIR_ABS if os.path.exists(OUTPUT_DIR_ABS) else str(BASE_DIR / "outputs" / "a1_ocr_mixed_text_table_GPT")
YAML_RULE_PATH = YAML_RULE_PATH_ABS if os.path.exists(YAML_RULE_PATH_ABS) else str(BASE_DIR / "configs" / "a1_text_only_rules.yaml")

# ========== Tiện ích chung ==========
def sha1_of_text(text: str) -> str:
    return hashlib.sha1((text or "").encode("utf-8")).hexdigest()

def detect_language_safe(text: str) -> str:
    try:
        return detect(text)
    except Exception:
        return "unknown"

def load_yaml_rules(path: str) -> dict:
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)
    return {}

def match_yaml_meta(file_name: str, rules: dict) -> dict:
    """Tìm meta mặc định dựa theo regex trong YAML."""
    if not rules or "files" not in rules:
        return {}
    for rule in rules["files"]:
        if re.search(rule.get("match", ""), file_name, flags=re.I):
            meta = rule.copy()
            meta.pop("match", None)
            return {**rules.get("defaults", {}), **meta}
    return rules.get("defaults", {})

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)

# ========== Helpers cho PDF scan ==========
def _rasterize_pdf_pages(pdf_path: Path, zoom: float = 2.0) -> List[Image.Image]:
    """Trả về list PIL.Image từ mỗi trang PDF (scan hoặc vector)."""
    imgs: List[Image.Image] = []
    with fitz.open(str(pdf_path)) as doc:
        mat = fitz.Matrix(zoom, zoom)  # ~144dpi nếu zoom=2.0
        for page in doc:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            mode = "RGB" if pix.n < 5 else "CMYK"
            img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
            imgs.append(img)
    return imgs

def _ocr_pil(pil_img: Image.Image, lang: str = OCR_LANG) -> str:
    """OCR trực tiếp PIL.Image bằng Tesseract (không cần bytes)."""
    return pytesseract.image_to_string(pil_img, lang=lang, config=OCR_CFG) or ""

def _save_png_debug(img: Image.Image, out_dir: Path, name: str):
    if not SAVE_PNG_DEBUG:
        return
    out = out_dir / f"{name}.png"
    try:
        img.save(out)
    except Exception:
        pass

# ========== Xử lý DOCX ==========
def read_docx_text_and_tables(file_path: Path) -> List[Dict[str, str]]:
    """Đọc nội dung từ DOCX, tách giữa đoạn văn và bảng."""
    doc = Document(str(file_path))
    results = []

    for block in doc.element.body:
        if block.tag.endswith("tbl"):  # Bảng
            rows = []
            for r in block.findall(".//w:tr", block.nsmap):
                cells = [c.text for c in r.findall(".//w:t", block.nsmap)]
                rows.append(" | ".join(cells))
            text_block = "\n".join(rows)
            results.append({"type": "table", "content": text_block})
        else:  # Đoạn văn
            text = "".join(t.text for t in block.findall(".//w:t", block.nsmap)).strip()
            if text:
                results.append({"type": "paragraph", "content": text})
    return results

# ========== Xử lý PDF (giữ nguyên + thêm fallback scan) ==========
def read_pdf_text_and_images(file_path: Path, out_dir: Optional[Path] = None,
                             page_start: Optional[int] = None,
                             page_end: Optional[int] = None) -> List[Dict[str, str]]:
    """
    Đọc PDF: text + ảnh (OCR ảnh khi cần).
    Fallback: nếu trang không có text/ảnh → rasterize toàn trang + OCR (PDF scan).
    Cho phép giới hạn phạm vi trang [page_start, page_end] (1-based, inclusive).
    """
    results: List[Dict[str, str]] = []

    # Mở fitz 1 lần (dùng cho fallback)
    try:
        fitz_doc = fitz.open(str(file_path))
    except Exception as e:
        print(f"⚠️ fitz.open lỗi {file_path}: {e}")
        fitz_doc = None

    # Tính phạm vi trang (1-based)
    try:
        total_pages = fitz_doc.page_count if fitz_doc else None
    except Exception:
        total_pages = None
    if page_start is None: page_start = 1
    if page_end   is None and total_pages: page_end = total_pages
    if page_end   is None: page_end = page_start
    # đảm bảo hợp lệ
    page_start = max(1, page_start)
    page_end   = max(page_start, page_end)

    try:
        with pdfplumber.open(str(file_path)) as pdf:
            # clamp theo số trang thực
            if total_pages:
                page_end = min(page_end, total_pages)

            for pi in range(page_start, page_end + 1):
                page = pdf.pages[pi - 1]
                page_had_content = False

                # 1) Text gốc
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""
                if text.strip():
                    results.append({"type": "paragraph", "content": f"[PAGE {pi}][TEXT]\n{text}"})
                    page_had_content = True

                # 2) OCR vùng ảnh nhúng
                try:
                    if getattr(page, "images", None):
                        for img in page.images:
                            x0, y0, x1, y1 = img["x0"], img["y0"], img["x1"], img["y1"]
                            pil_img = page.within_bbox((x0, y0, x1, y1)).to_image(resolution=300).original
                            if not isinstance(pil_img, Image.Image):
                                pil_img = Image.fromarray(pil_img)
                            pil_img = pil_img.convert("RGB")
                            ocr_txt = _ocr_pil(pil_img, lang=OCR_LANG)
                            if ocr_txt.strip():
                                results.append({"type": "table", "content": f"[PAGE {pi}][OCR_IMG]\n{ocr_txt}"})
                                page_had_content = True
                except Exception:
                    pass

                # 3) Fallback toàn trang nếu rỗng
                if not page_had_content and fitz_doc is not None:
                    try:
                        pg = fitz_doc.load_page(pi - 1)
                        pix = pg.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                        mode = "RGB" if pix.n < 5 else "CMYK"
                        pil_full = Image.frombytes(mode, (pix.width, pix.height), pix.samples).convert("RGB")
                        ocr_full = _ocr_pil(pil_full, lang=OCR_LANG)
                        if ocr_full.strip():
                            results.append({"type": "paragraph", "content": f"[PAGE {pi}][OCR_PAGE]\n{ocr_full}"})
                    except Exception as e:
                        print(f"⚠️ PDF OCR fallback error (page {pi}): {e}")

    except Exception as e:
        # pdfplumber hỏng → rasterize theo phạm vi
        print(f"⚠️ PDF đọc lỗi {file_path}: {e}")
        if fitz_doc is not None:
            try:
                if total_pages:
                    page_end = min(page_end, total_pages)
                for pi in range(page_start, page_end + 1):
                    pg = fitz_doc.load_page(pi - 1)
                    pix = pg.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                    mode = "RGB" if pix.n < 5 else "CMYK"
                    pil_full = Image.frombytes(mode, (pix.width, pix.height), pix.samples).convert("RGB")
                    ocr_txt = _ocr_pil(pil_full, lang=OCR_LANG)
                    if ocr_txt.strip():
                        results.append({"type": "paragraph", "content": f"[PAGE {pi}][OCR_PAGE]\n{ocr_txt}"})
            except Exception as e2:
                print(f"⚠️ PDF rasterize error {file_path}: {e2}")

    try:
        if fitz_doc is not None:
            fitz_doc.close()
    except Exception:
        pass

    return results



# ========== Xử lý hình ảnh ==========
def read_docx_text_and_tables(file_path: Path) -> List[Dict[str, str]]:
    """
    Đọc DOCX bằng API python-docx (tránh XML thô).
    - Đoạn văn: p.text
    - Bảng: nối cell theo hàng, ' | ' giữa cell
    """
    doc = Document(str(file_path))
    results: List[Dict[str, str]] = []

    # paragraphs
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            results.append({"type": "paragraph", "content": txt})

    # tables
    for tbl in doc.tables:
        rows_text = []
        for row in tbl.rows:
            cells = []
            for cell in row.cells:
                cells.append((cell.text or "").replace("\n", " ").strip())
            rows_text.append(" | ".join(cells))
        if rows_text:
            results.append({"type": "table", "content": "\n".join(rows_text)})

    return results

# ========== Xử lý TXT ==========
def read_txt(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

# ========== Xử lý Excel (multi-sheet) ==========
def read_excel_as_text(file_path: Path) -> List[Dict[str, str]]:
    """
    Đọc file Excel (.xlsx, .xls) gồm nhiều sheet.
    Mỗi sheet trả về 1 block {"sheet": <tên>, "content": <text bảng>}.
    """
    results = []
    try:
        sheets = pd.read_excel(str(file_path), sheet_name=None, header=None)
        for sheet_name, df in sheets.items():
            df = df.fillna("")
            rows = []
            for row in df.values.tolist():
                row_text = " | ".join(str(cell).strip() for cell in row)
                rows.append(row_text)
            table_text = "\n".join(rows)
            results.append({"sheet": sheet_name, "content": table_text})
    except Exception as e:
        print(f"⚠️ Excel đọc lỗi {file_path}: {e}")
    return results


def _post_clean(text: str) -> str:
    """
    Clean nhẹ:
    - Gộp chữ số bị tách bởi khoảng trắng (12 345 -> 12345)
    - Xoá '|' lẻ ở cuối dòng
    - Chuẩn hoá khoảng trắng thừa
    """
    if not text:
        return text
    text = re.sub(r'(?<=\d)\s+(?=\d)', '', text)              # gộp số bị chẻ
    text = re.sub(r'\|\s*$', '', text, flags=re.MULTILINE)    # xóa '|' cuối dòng
    text = re.sub(r'[ \t]{2,}', ' ', text)                    # thu gọn nhiều space/tab
    return text


def ocr_image_to_text(file_path: Path) -> str:
    """OCR ảnh sang text (ổn định) — luôn ép RGB trước khi đưa vào Tesseract."""
    img = Image.open(str(file_path)).convert("RGB")
    return pytesseract.image_to_string(img, lang=OCR_LANG, config=OCR_CFG)


# ========== Xử lý chính ==========
def process_file(file_path: Path, yaml_rules: dict):
    file_name = file_path.stem
    output_dir = Path(OUTPUT_DIR)
    ensure_dir(output_dir)

    meta = {
        "file": file_name,
        "source_path": str(file_path.resolve()),
        "ocr_lang": OCR_LANG,
        "ocr_cfg": OCR_CFG,
    }
    meta.update(match_yaml_meta(file_name, yaml_rules))

    ext = file_path.suffix.lower()
    has_table = False
    gpt_applied = False
    combined_text = ""

    try:
        if ext == ".docx":
            blocks = read_docx_text_and_tables(file_path)
            for block in blocks:
                if block["type"] == "table":
                    has_table = True
                    gpt_applied = True
                    new_txt = _enhance_with_gpt(block["content"], meta, None)
                    combined_text += "\n" + new_txt
                else:
                    combined_text += "\n" + block["content"]

        elif ext == ".pdf":
            # Nếu PAGE_MODE=1 → hiểu start/end là phạm vi trang
            PAGE_MODE = os.getenv("PAGE_MODE", "0") == "1"
            page_start = None
            page_end = None
            if PAGE_MODE and "_page_start" in meta and "_page_end" in meta:
                page_start = meta["_page_start"]
                page_end   = meta["_page_end"]

            blocks = read_pdf_text_and_images(file_path, out_dir=output_dir,
                                              page_start=page_start, page_end=page_end)
            for block in blocks:
                if block["type"] == "table":
                    has_table = True
                    gpt_applied = True
                    new_txt = _enhance_with_gpt(block["content"], meta, None)
                    combined_text += "\n" + new_txt
                else:
                    combined_text += "\n" + block["content"]


        elif ext in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"]:
            has_table = True
            gpt_applied = True
            txt = ocr_image_to_text(file_path)
            combined_text += "\n" + _enhance_with_gpt(txt, meta, str(file_path))

        elif ext in [".txt", ".csv"]:
            txt = read_txt(file_path)
            combined_text += "\n" + txt

        elif ext in [".xlsx", ".xls"]:
            excel_blocks = read_excel_as_text(file_path)
            if excel_blocks:
                has_table = True
                gpt_applied = True
                all_text = "\n\n".join([f"--- Sheet: {b['sheet']} ---\n{b['content']}" for b in excel_blocks])
                enhanced_txt = _enhance_with_gpt(all_text, meta, None)
                combined_text += "\n" + enhanced_txt

        else:
            print(f"⚠️ Không hỗ trợ định dạng {ext}")
            return



        # Clean nhẹ để tăng chất lượng truy vấn số/bảng
        combined_text = _post_clean(combined_text)

        # Nếu vẫn rỗng → thử fallback lần cuối cho PDF/Ảnh (zoom cao hơn)
        if not combined_text.strip():
            if ext == ".pdf":
                print(f"⚠️ {file_name}: combined_text trống → fallback rasterize zoom=3.0")
                try:
                    with fitz.open(str(file_path)) as _doc:
                        chunks = []
                        for pi in range(_doc.page_count):
                            pg = _doc.load_page(pi)
                            pix = pg.get_pixmap(matrix=fitz.Matrix(3.0, 3.0), alpha=False)
                            mode = "RGB" if pix.n < 5 else "CMYK"
                            pil_full = Image.frombytes(mode, (pix.width, pix.height), pix.samples).convert("RGB")
                            ocr_full = _ocr_pil(pil_full, lang=OCR_LANG)
                            if ocr_full.strip():
                                chunks.append(f"[PAGE {pi+1}][OCR_PAGE]\n{ocr_full}")
                        combined_text = "\n\n".join(chunks)
                except Exception as _e:
                    print(f"⚠️ Fallback PDF zoom=3.0 lỗi: {_e}")

            elif ext in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"]:
                print(f"⚠️ {file_name}: OCR ảnh rỗng → thử lại convert RGB & psm=6")
                try:
                    combined_text = ocr_image_to_text(file_path)
                except Exception as _e:
                    print(f"⚠️ Fallback OCR ảnh lỗi: {_e}")

        # enrich meta
        lang = detect_language_safe(combined_text)
        meta.update({
            "language": lang,
            "has_table": has_table,
            "gpt_applied": gpt_applied,
            "text_sha1": sha1_of_text(combined_text),
            "empty_after_fallback": not bool(combined_text.strip())
        })

        # Nếu vẫn trống hoàn toàn → ghi marker để tránh file 0 KB
        if not combined_text.strip():
            combined_text = "[EMPTY] No text extracted by OCR/PDF. See meta.empty_after_fallback=true"

        # Ghi file output
        txt_out = output_dir / f"{file_name}_text.txt"
        meta_out = output_dir / f"{file_name}_meta.json"
        with open(txt_out, "w", encoding="utf-8") as f:
            f.write(combined_text.strip())
        with open(meta_out, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

        print(f"✅ Saved: {txt_out.name}, {meta_out.name}  | len={len(combined_text)}")

    except Exception as e:
        print(f"❌ Lỗi xử lý {file_path}: {e}")

# ========== CLI ==========
def main():
    parser = argparse.ArgumentParser(
        description="OCR + GPT cho tài liệu hỗn hợp (text + bảng + ảnh)"
    )
    parser.add_argument("--clean", choices=["y", "n"], default="n")
    parser.add_argument("--start", type=int, default=1, help="Số thứ tự file bắt đầu (tính từ 1) hoặc TRANG bắt đầu nếu PAGE_MODE=1")
    parser.add_argument("--end", type=int, default=None, help="Số thứ tự file kết thúc (hoặc TRANG kết thúc nếu PAGE_MODE=1)")
    parser.add_argument("--file_index", type=int, default=1, help="(Chỉ khi PAGE_MODE=1) chọn file thứ mấy trong INPUT để cắt trang (mặc định 1)")
    args = parser.parse_args()

    yaml_rules = load_yaml_rules(YAML_RULE_PATH)
    files = sorted(list(Path(INPUT_DIR).rglob("*.*")))

    PAGE_MODE = os.getenv("PAGE_MODE", "0") == "1"

    if PAGE_MODE:
        # Hiểu --start/--end là trang
        # Chọn file PDF theo --file_index (1-based)
        pdfs = [f for f in files if f.suffix.lower() == ".pdf"]
        if not pdfs:
            print("⚠️ PAGE_MODE=1 nhưng không tìm thấy PDF nào trong INPUT_DIR.")
            return
        file_idx = max(1, min(args.file_index, len(pdfs))) - 1
        target = pdfs[file_idx]

        # Lưu thông tin trang vào meta thông qua biến tạm
        global _PAGE_RANGE_CACHE
        _PAGE_RANGE_CACHE = (args.start, args.end)

        print("=== CẤU HÌNH HIỆN TẠI (PAGE_MODE=1) ===")
        print(f"📂 INPUT_DIR   : {INPUT_DIR}")
        print(f"📦 OUTPUT_DIR  : {OUTPUT_DIR}")
        print(f"⚙️ YAML_RULE   : {YAML_RULE_PATH}")
        print(f"🎯 FILE        : {target.name}")
        print(f"📄 PAGES       : {args.start} → {args.end}")
        print("==========================")

        # Gọi process_file chỉ cho 1 file PDF, và bơm _page_start/_page_end vào meta
        def _process_pdf_with_pages(fpath: Path, rules: dict):
            # bơm vào meta qua match_yaml_meta → cách nhanh là set vào rules.defaults tạm thời
            # tránh side-effects: copy rules
            rules = dict(rules) if rules else {}
            defaults = dict(rules.get("defaults", {}))
            defaults["_page_start"] = args.start
            defaults["_page_end"] = args.end if args.end else args.start
            rules["defaults"] = defaults
            process_file(fpath, rules)

        _process_pdf_with_pages(Path(target), yaml_rules)

    else:
        # Hành vi cũ: --start/--end là chỉ số file
        selected = files[args.start - 1 : args.end] if args.end else files[args.start - 1 :]

        print("=== CẤU HÌNH HIỆN TẠI ===")
        print(f"📂 INPUT_DIR   : {INPUT_DIR}")
        print(f"📦 OUTPUT_DIR  : {OUTPUT_DIR}")
        print(f"⚙️ YAML_RULE   : {YAML_RULE_PATH}")
        print(f"🧠 Tổng số file: {len(files)} | Sẽ xử lý: {len(selected)} (từ {args.start} đến {args.end or len(files)})")
        print("==========================")

        for file_path in tqdm(selected, desc="Processing"):
            process_file(Path(file_path), yaml_rules)

    print("🎯 Hoàn tất A2. Kiểm tra thư mục output để xem kết quả.")


if __name__ == "__main__":
    main()
