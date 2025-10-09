# -*- coding: utf-8 -*-
"""
a1_ocr_text_only_v1.py — IO đa định dạng + OCR (TSV reflow) + YAML rules + anchors
Giữ nguyên default paths & YAML của bạn. PDF/IMG xuất per-page; DOCX xuất 1 cặp file có anchors.
"""

from __future__ import annotations
import os, re, glob, json, argparse, hashlib, shutil, time
from typing import Optional, Tuple, Dict, List
from pathlib import Path

import numpy as np
import cv2
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from pytesseract import Output as TessOutput
import pandas as pd
from tqdm import tqdm

# YAML
try:
    import yaml  # type: ignore
except Exception:
    yaml = None
    try:
        from ruamel import yaml as ruamel_yaml  # type: ignore
    except Exception:
        ruamel_yaml = None

# DOCX
try:
    import docx
except ImportError:
    docx = None

# ============== Defaults (GIỮ NGUYÊN NHƯ FILE CỦA BẠN) ==============
INPUT_DIR_DEFAULT  = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\a_text_only_inputs_test"
OUTPUT_DIR_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\a_text_only_outputs2"
RULES_PATH_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_text_only_rules.yaml"

OCR_LANG_DEFAULT = "vie+eng"
OCR_CFG_DEFAULT  = "--psm 6 preserve_interword_spaces=1"
APPEND_MODE = False

TESSERACT_CMD = os.environ.get("TESSERACT_CMD")
if TESSERACT_CMD and os.path.isfile(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# ================= Common utils =================
def ensure_dir(p: str) -> None:
    os.makedirs(p, exist_ok=True)

def sha1_text(s: str) -> str:
    return hashlib.sha1((s or "").encode("utf-8")).hexdigest()

def sha1_file(path: str) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def clean_txt(s: str) -> str:
    if not s: return ""
    s = re.sub(r"[|¦•∙·■□▪◦●○◻◼▶►◆◇★☆■□]{2,}", " ", s)
    s = s.replace("\xa0", " ")
    s = re.sub(r"[^\S\r\n]{2,}", " ", s)
    s = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", s)
    s = re.sub(r"(?<=\d)(?=[A-Za-z])", " ", s)
    s = re.sub(r"(?<=[a-z])(?=\d)", " ", s)
    return s.strip()

# language detection: optional langdetect
def detect_language(text: str) -> str:
    if not text: return "unknown"
    try:
        from langdetect import detect  # type: ignore
        lang = detect(text[:5000])
        return lang or "unknown"
    except Exception:
        vi_marks = re.findall(r"[ăâêôơưđáàảãạéèẻẽẹíìỉĩịóòỏõọúùủũụýỳỷỹỵ]", text.lower())
        if len(vi_marks) >= 3: return "vi"
        if re.search(r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\b", text, re.I):
            return "en"
        return "unknown"

# ================= YAML rules =================
def load_rules(path: str) -> dict:
    if not path or not os.path.exists(path): return {}
    try:
        if yaml:
            with open(path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        elif ruamel_yaml:
            with open(path, "r", encoding="utf-8") as f:
                return ruamel_yaml.YAML(typ="safe").load(f) or {}
    except Exception as e:
        print(f"⚠️ Không đọc được YAML rules: {e}")
    return {}

def match_rules(file_path: str, text: str, rules: dict) -> dict:
    if not rules: return {}
    res = {}
    defaults = rules.get("defaults", {}) or {}
    files = rules.get("files", []) or []
    fname = os.path.basename(file_path)
    for item in files:
        pat = item.get("match")
        if not pat: continue
        try:
            if re.search(pat, fname, flags=re.IGNORECASE):
                res = {**defaults, **{k:v for k,v in item.items() if k!="match"}}
                break
        except re.error:
            continue
    if not res:
        for item in files:
            tpat = item.get("text_match")
            if not tpat: continue
            try:
                if re.search(tpat, text or "", flags=re.IGNORECASE):
                    res = {**defaults, **{k:v for k,v in item.items() if k not in ("match","text_match")}}
                    break
            except re.error:
                continue
    return res or defaults

# ================= Anchors =================
RE_CHAPTER = re.compile(r"^\s*(CHAPTER|CHAP\.?)\s+([IVXLC]+)\b(.*)$", re.I)
RE_SECTION = re.compile(r"^\s*(SECTION|SEC\.?)\s+([A-Z0-9]+)\b(.*)$", re.I)
RE_PART    = re.compile(r"^\s*(PART|PT\.?)\s+([A-Z0-9]+)\b(.*)$", re.I)

def add_anchors(text: str) -> Tuple[str, List[str]]:
    anchors: List[str] = []
    out = []
    for line in (text or "").splitlines():
        raw = line.strip()
        if not raw:
            out.append("")
            continue
        m = RE_CHAPTER.match(raw)
        if m:
            head = f"CHAPTER {m.group(2)}{m.group(3)}".strip()
            anchors.append(head)
            out.append(f"# {head}")
            continue
        m = RE_SECTION.match(raw)
        if m:
            head = f"SECTION {m.group(2)}{m.group(3)}".strip()
            anchors.append(head)
            out.append(f"## {head}")
            continue
        m = RE_PART.match(raw)
        if m:
            head = f"PART {m.group(2)}{m.group(3)}".strip()
            anchors.append(head)
            out.append(f"### {head}")
            continue
        out.append(re.sub(r"^\s*[-•·]\s*", "- ", raw))
    txt2 = "\n".join(out)
    txt2 = re.sub(r"\n{3,}", "\n\n", txt2).strip()
    # unique, stable order
    anchors = list(dict.fromkeys(anchors))
    return txt2, anchors

# ================= TSV reflow OCR =================
def reflow_from_tsv(tsv: Dict[str, List]) -> str:
    n = len(tsv.get("text", []))
    lines = []
    groups: Dict[Tuple[int,int,int], List[int]] = {}
    for i in range(n):
        t = (tsv["text"][i] or "").strip()
        if not t: continue
        key = (int(tsv["block_num"][i]), int(tsv["par_num"][i]), int(tsv["line_num"][i]))
        groups.setdefault(key, []).append(i)
    for _, idxs in sorted(groups.items()):
        idxs = sorted(idxs, key=lambda k: int(tsv["left"][k]))
        txt = " ".join((tsv["text"][k] or "").strip() for k in idxs).strip()
        if not txt: continue
        y = int(min(int(tsv["top"][k]) for k in idxs))
        lines.append((y, txt))
    lines.sort(key=lambda r: r[0])
    out = []
    for _, s in lines:
        s = re.sub(r"(?<!^)\s(?=(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\.?\b)", "\n", s)
        s = re.sub(r"(?<!^)\s(?=\d{1,3}(?:[.,]\d{3}){2,}\b)", "\n", s)
        out.extend([p.strip() for p in s.split("\n") if p.strip()])
    return "\n".join(out)

def ocr_img_tsv(img_bgr, ocr_lang: str, ocr_cfg: str) -> str:
    try:
        # dùng đúng ocr_cfg người dùng truyền
        tsv = pytesseract.image_to_data(img_bgr, lang=ocr_lang, config=(ocr_cfg or "").strip(), output_type=TessOutput.DICT)
        return clean_txt(reflow_from_tsv(tsv))
    except Exception as e:
        print(f"⚠️ OCR TSV error: {e}")
        return ""

# ================= Readers =================
def read_pdf_per_page(pdf_path: str, dpi: int, ocr_lang: str, ocr_cfg: str,
                      start_page: Optional[int], end_page: Optional[int]) -> List[Tuple[int,str]]:
    texts = []
    try:
        pages = convert_from_path(pdf_path, dpi=dpi) if (start_page is None and end_page is None) \
            else convert_from_path(pdf_path, dpi=dpi, first_page=start_page or 1, last_page=end_page or (start_page or 1))
    except Exception as e:
        print(f"⚠️ PDF open error {pdf_path}: {e}")
        return texts
    s = start_page or 1
    total = len(pages)
    for idx, page in enumerate(pages, start=0):
        page_no = s + idx
        img = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
        txt = ocr_img_tsv(img, ocr_lang, ocr_cfg)
        texts.append((page_no, txt))
    return texts

def read_docx_flat(path: str) -> str:
    if docx is None:
        raise ImportError("⚠️ Cần python-docx để đọc DOCX.")
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in d.tables:
        parts.append("========== TABLE ==========")
        for r in tbl.rows:
            cells = [clean_txt(c.text) for c in r.cells]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()

def read_excel_csv(path: str, mode: str="summary") -> str:
    ext = Path(path).suffix.lower()
    texts = []
    try:
        if ext == ".csv":
            dfs = {"CSV": pd.read_csv(path, dtype=str)}
        else:
            xls = pd.ExcelFile(path, engine="openpyxl")
            dfs = {sh: pd.read_excel(xls, sheet_name=sh, dtype=str) for sh in xls.sheet_names}
        for sheet_name, df in dfs.items():
            if df.empty: continue
            df = df.fillna("").astype(str)
            if mode == "raw":
                headers = list(df.columns)
                rows = [" | ".join(row.values) for _, row in df.iterrows()]
                sheet_text = f"\n\n========== SHEET: {sheet_name.upper()} ==========\n" + " | ".join(headers) + "\n" + "\n".join(rows)
            else:
                keep = df.loc[:, ~df.columns.str.match(r"^Unnamed($|:)")]
                headers = list(keep.columns)
                sheet_text = [f"\n========== SHEET: {sheet_name.upper()} ==========",
                              f"Các cột gồm: {', '.join(headers)}."]
                for _, row in keep.iterrows():
                    if not any(v.strip() for v in row.values): continue
                    pairs = [f"{h}: {v}" for h, v in row.items() if v.strip()]
                    sheet_text.append("; ".join(pairs))
                sheet_text = "\n".join(sheet_text)
            texts.append(sheet_text.strip())
        return "\n\n".join(texts).strip()
    except Exception as e:
        print(f"⚠️ Lỗi đọc {path}: {e}")
        return ""

# ================= Save =================
def save_text_meta(text: str, meta: dict, out_txt: str, out_meta: str):
    ensure_dir(os.path.dirname(out_txt))
    if APPEND_MODE and os.path.exists(out_txt) and os.path.exists(out_meta):
        try:
            with open(out_txt, "r", encoding="utf-8") as f:
                if sha1_text(f.read()) == sha1_text(text):
                    print(f"⏭️ Bỏ qua (không đổi): {os.path.basename(out_txt)}")
                    return
        except Exception:
            pass
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(text)
    with open(out_meta, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    print(f"📝 Saved: {os.path.basename(out_txt)}, {os.path.basename(out_meta)}")

# ================= Process =================
def enrich_with_rules_and_anchors(file_path: str, text: str, meta: dict, rules: dict) -> Tuple[str, dict]:
    text_clean = clean_txt(text)
    text_anchor, anchors = add_anchors(text_clean)
    if anchors: meta["anchors"] = anchors
    matched = match_rules(file_path, text_anchor, rules)
    if matched:
        meta.update({
            "company": matched.get("company"),
            "doc_title": matched.get("doc_title"),
            "doc_type": matched.get("doc_type"),
            "department": matched.get("department"),
            "lob": matched.get("lob"),
            "appendix_id": matched.get("appendix_id"),
            "effective_date": matched.get("effective_date"),
            "version": matched.get("version"),
            "confidentiality": matched.get("confidentiality"),
        })
    return text_anchor, meta

def process_file(file_path: str, input_root: str, output_root: str,
                 ocr_lang: str, ocr_cfg: str, rules: dict,
                 dpi: int, start_page: Optional[int], end_page: Optional[int],
                 excel_mode: str):

    rel = os.path.relpath(file_path, input_root)
    base = Path(file_path).stem
    out_dir = os.path.join(output_root, os.path.dirname(rel))
    ensure_dir(out_dir)

    ext = Path(file_path).suffix.lower()
    src_sha = sha1_file(file_path)

    if ext in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"]:
        img = cv2.cvtColor(np.array(Image.open(file_path).convert("RGB")), cv2.COLOR_RGB2BGR)
        txt = ocr_img_tsv(img, ocr_lang, ocr_cfg)
        out_txt  = os.path.join(out_dir, f"{base}_page1_text.txt")
        out_meta = os.path.join(out_dir, f"{base}_page1_meta.json")
        meta = {
            "file": base, "page": 1,
            "source_path": os.path.abspath(file_path), "source_doc_sha1": src_sha,
            "language": detect_language(txt),
            "ocr_lang": ocr_lang, "ocr_cfg": ocr_cfg,
            "text_sha1": sha1_text(txt),
        }
        text2, meta2 = enrich_with_rules_and_anchors(file_path, txt, meta, rules)
        save_text_meta(text2, meta2, out_txt, out_meta)
        return

    if ext == ".pdf":
        pages_text = read_pdf_per_page(file_path, dpi, ocr_lang, ocr_cfg, start_page, end_page)
        print(f"📄 {os.path.basename(file_path)} → {len(pages_text)} trang")
        for pno, txt in pages_text:
            out_txt  = os.path.join(out_dir, f"{base}_page{pno}_text.txt")
            out_meta = os.path.join(out_dir, f"{base}_page{pno}_meta.json")
            meta = {
                "file": base, "page": pno,
                "source_path": os.path.abspath(file_path), "source_doc_sha1": src_sha,
                "language": detect_language(txt),
                "ocr_lang": ocr_lang, "ocr_cfg": ocr_cfg,
                "text_sha1": sha1_text(txt),
            }
            text2, meta2 = enrich_with_rules_and_anchors(file_path, txt, meta, rules)
            save_text_meta(text2, meta2, out_txt, out_meta)
        return

    if ext in [".doc", ".docx"]:
        txt = read_docx_flat(file_path)
        out_txt  = os.path.join(out_dir, f"{base}_page1_text.txt")
        out_meta = os.path.join(out_dir, f"{base}_page1_meta.json")
        meta = {
            "file": base, "page": 1, "pagination": "none-docx",
            "source_path": os.path.abspath(file_path), "source_doc_sha1": src_sha,
            "language": detect_language(txt),
            "text_sha1": sha1_text(txt),
        }
        # KHÔNG ghi ocr_* vì không OCR
        text2, meta2 = enrich_with_rules_and_anchors(file_path, txt, meta, rules)
        save_text_meta(text2, meta2, out_txt, out_meta)
        return

    if ext in [".xls", ".xlsx", ".csv"]:
        txt = read_excel_csv(file_path, mode=excel_mode)
        out_txt  = os.path.join(out_dir, f"{base}_page1_text.txt")
        out_meta = os.path.join(out_dir, f"{base}_page1_meta.json")
        meta = {
            "file": base, "page": 1, "pagination": "none-table",
            "source_path": os.path.abspath(file_path), "source_doc_sha1": src_sha,
            "language": detect_language(txt),
            "text_sha1": sha1_text(txt),
        }
        text2, meta2 = enrich_with_rules_and_anchors(file_path, txt, meta, rules)
        save_text_meta(text2, meta2, out_txt, out_meta)
        return

    if ext == ".txt":
        txt = open(file_path, "r", encoding="utf-8", errors="ignore").read()
        out_txt  = os.path.join(out_dir, f"{base}_page1_text.txt")
        out_meta = os.path.join(out_dir, f"{base}_page1_meta.json")
        meta = {
            "file": base, "page": 1, "pagination": "none-plain",
            "source_path": os.path.abspath(file_path), "source_doc_sha1": src_sha,
            "language": detect_language(txt),
            "text_sha1": sha1_text(txt),
        }
        text2, meta2 = enrich_with_rules_and_anchors(file_path, txt, meta, rules)
        save_text_meta(text2, meta2, out_txt, out_meta)
        return

    print(f"⚠️ Bỏ qua (không hỗ trợ): {file_path}")

# ================= Main =================
def main():
    parser = argparse.ArgumentParser("A1 v1 — per-page PDF/IMG, DOCX flat + anchors, YAML rules")
    parser.add_argument("--input", type=str, default=INPUT_DIR_DEFAULT)
    parser.add_argument("--out", type=str, default=OUTPUT_DIR_DEFAULT)
    parser.add_argument("--rules", type=str, default=RULES_PATH_DEFAULT)
    parser.add_argument("--ocr-lang", type=str, default=OCR_LANG_DEFAULT)
    parser.add_argument("--ocr-cfg", type=str, default=OCR_CFG_DEFAULT)
    parser.add_argument("--dpi", type=int, default=400)
    parser.add_argument("--clean", choices=["y","a","n","ask"], default="ask")
    parser.add_argument("--start", type=int, default=None)
    parser.add_argument("--end", type=int, default=None)
    parser.add_argument("--excel-mode", choices=["raw","summary"], default="summary")
    args = parser.parse_args()

    if os.path.exists(args.out):
        choice = args.clean
        if choice == "ask":
            try:
                choice = input(f"⚠️ Output '{args.out}' đã tồn tại. y=xoá, a=append, n=bỏ qua: ").strip().lower()
            except Exception:
                choice = "a"
        if choice == "y":
            shutil.rmtree(args.out, ignore_errors=True); print(f"🗑️ Đã xoá {args.out}")
        elif choice == "a":
            global APPEND_MODE; APPEND_MODE = True
            print("➕ APPEND_MODE = True (bỏ qua file không đổi)")
        elif choice == "n":
            print("⏭️ Bỏ qua toàn bộ."); return
        else:
            print("❌ Lựa chọn không hợp lệ."); return

    ensure_dir(args.out)
    files = glob.glob(os.path.join(args.input, "**", "*.*"), recursive=True)
    if not files:
        print("⚠️ Không tìm thấy file nào."); return

    rules = load_rules(args.rules)
    if rules: print(f"🧩 Loaded rules: {args.rules}")
    else:     print(f"ℹ️ Không có rules (vẫn chạy).")

    print(f"📂 Input : {args.input}")
    print(f"📦 Output: {args.out}")
    print(f"🧮 Tổng số file: {len(files)}")
    print(f"🧭 PDF pages: {args.start or 1} → {args.end or 'all'}")

    for f in files:
        process_file(f, args.input, args.out, args.ocr_lang, args.ocr_cfg, rules,
                     dpi=args.dpi, start_page=args.start, end_page=args.end,
                     excel_mode=args.excel_mode)

    print("\n✅ Done.")

if __name__ == "__main__":
    main()
