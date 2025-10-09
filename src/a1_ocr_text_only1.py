# -*- coding: utf-8 -*-
"""
a1_ocr_text_only.py — OCR và đọc văn bản đa định dạng (PDF, DOCX, EXCEL, CSV, IMAGE)

Mục tiêu:
---------
- Đọc tất cả các loại tài liệu (PDF, DOCX, XLSX, CSV, IMAGE, TXT)
- Không sinh ra bất kỳ file CSV/Excel nào.
- Chuyển toàn bộ nội dung (kể cả bảng scan, sơ đồ) sang text duy nhất.
- Xuất đúng 2 file/trang:
    <base>_page{n}_text.txt
    <base>_page{n}_meta.json
- Giữ cấu trúc thư mục mirror từ inputs sang outputs.
- Khi OCR ảnh hoặc PDF scan → dùng TSV reflow để tái cấu trúc dòng/bảng.
- Bổ sung metadata mở rộng từ YAML rule và chèn heading anchors hỗ trợ chunking RAG.

Đường dẫn mặc định:
-------------------
Input : D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\a_text_only_inputs_test
Output: D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\a_text_only_outputs
Rules : D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_text_only_rules.yaml

Yêu cầu thư viện:
-----------------
pip install pdf2image pillow opencv-python-headless numpy pytesseract python-docx pandas openpyxl tqdm pyyaml
và cài đặt Tesseract (tesseract.exe có trong PATH hoặc đặt env TESSERACT_CMD)
"""

from __future__ import annotations
import os, re, glob, json, argparse, hashlib, shutil, time
from typing import Optional, Tuple, Dict, List
from pathlib import Path

import numpy as np
import cv2
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from pytesseract import Output as TessOutput
import pandas as pd
from tqdm import tqdm

# YAML loader (ưu tiên PyYAML, fallback ruamel)
try:
    import yaml  # type: ignore
except Exception:
    yaml = None
    try:
        from ruamel import yaml as ruamel_yaml  # type: ignore
    except Exception:
        ruamel_yaml = None

try:
    import docx
except ImportError:
    docx = None

# =========================
# ⚙️ CẤU HÌNH CƠ BẢN
# =========================
INPUT_DIR_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\inputs\a_text_only_inputs_test"
OUTPUT_DIR_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\outputs\a_text_only_outputs1"
RULES_PATH_DEFAULT = r"D:\1.TLAT\3. ChatBot_project\1_Insurance_Strategy\configs\a1_text_only_rules.yaml"

OCR_LANG_DEFAULT = "vie+eng"
OCR_CFG_DEFAULT  = "--psm 6 preserve_interword_spaces=1"
APPEND_MODE = False

TESSERACT_CMD = os.environ.get("TESSERACT_CMD", None)
if TESSERACT_CMD and os.path.isfile(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# =========================
# ⚙️ HỖ TRỢ
# =========================
def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)

def clean_txt_chars(s: str) -> str:
    """Chuẩn hoá văn bản OCR: loại bullet rác & khoảng trắng thừa."""
    if not s: return ""
    s = re.sub(r"[|¦•∙·■□▪◦●○◻◼▶►•●◆◇★☆■□-]{2,}", " ", s)
    s = re.sub(r"\xa0", " ", s)             # non-breaking space
    s = re.sub(r"[^\S\r\n]{2,}", " ", s)    # nhiều space -> 1
    s = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", s)
    # fix dính chữ thường gặp khi OCR từ DOCX/PDF
    s = re.sub(r"(?<=\d)(?=[A-Za-z])", " ", s)
    s = re.sub(r"(?<=[a-z])(?=\d)", " ", s)
    return s.strip()

def _sha1_text(s: str) -> str:
    return hashlib.sha1((s or "").encode("utf-8")).hexdigest()

def _sha1_file(path: str) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def detect_language(text: str) -> str:
    if not text: return "vi"
    vi_marks = re.findall(r"[ăâêôơưđáàảãạéèẻẽẹíìỉĩịóòỏõọúùủũụýỳỷỹỵ]", text.lower())
    if len(vi_marks) >= 3: return "vi"
    if re.search(r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\b", text, re.IGNORECASE):
        return "en"
    return "vi"

# =========================
# ⚙️ ĐỌC YAML RULES
# =========================
def load_rules(path: str) -> dict:
    if not path or not os.path.exists(path):
        return {}
    try:
        if yaml:
            with open(path, "r", encoding="utf-8") as f:
                return yaml.safe_load(f) or {}
        elif ruamel_yaml:
            with open(path, "r", encoding="utf-8") as f:
                return ruamel_yaml.YAML(typ="safe").load(f) or {}
    except Exception as e:
        print(f"⚠️ Không đọc được YAML rules: {e}")
    return {}

def match_rules(file_path: str, text: str, rules: dict) -> dict:
    """
    Tìm rule khớp theo tên file (ưu tiên) hoặc pattern trong text.
    YAML format gợi ý:
    ---
    files:
      - match: "(?i)Appen6.*Retention"
        company: "UIC"
        doc_title: "APPENDIX 6 - RETENTION GUIDELINE"
        doc_type: "Underwriting Guideline / Appendix"
        appendix_id: "Appendix 6"
    defaults:
      department: "UWRI"
      confidentiality: "Internal"
    """
    if not rules: return {}
    res = {}
    defaults = rules.get("defaults", {}) or {}
    files = rules.get("files", []) or []
    fname = os.path.basename(file_path)
    for item in files:
        pat = item.get("match")
        if not pat: continue
        try:
            if re.search(pat, fname, flags=re.IGNORECASE):
                res = {**defaults, **{k:v for k,v in item.items() if k!="match"}}
                break
        except re.error:
            continue

    # Nếu chưa khớp theo tên file, thử khớp theo text
    if not res and files:
        for item in files:
            txt_pat = item.get("text_match")
            if not txt_pat: continue
            try:
                if re.search(txt_pat, text or "", flags=re.IGNORECASE):
                    res = {**defaults, **{k:v for k,v in item.items() if k not in ("match","text_match")}}
                    break
            except re.error:
                continue
    # Bồi defaults nếu có
    if defaults and res:
        res = {**defaults, **res}
    return res

# =========================
# ⚙️ TSV REFLOW cho OCR bảng/scan
# =========================
def reflow_lines_from_tsv_dict(data: Dict[str, List], y_tol: int = 4) -> str:
    n = len(data.get("text", []))
    groups: Dict[Tuple[int,int,int], List[int]] = {}
    for i in range(n):
        t = (data["text"][i] or "").strip()
        if not t: continue
        key = (int(data["block_num"][i]), int(data["par_num"][i]), int(data["line_num"][i]))
        groups.setdefault(key, []).append(i)

    lines = []
    for key, idxs in sorted(groups.items()):
        idxs = sorted(idxs, key=lambda k: int(data["left"][k]))
        txt = " ".join((data["text"][k] or "").strip() for k in idxs).strip()
        if not txt: continue
        y = int(min(int(data["top"][k]) for k in idxs))
        lines.append({"y": y, "text": txt})

    lines.sort(key=lambda r: r["y"])
    out_lines = []
    for ln in lines:
        s = ln["text"]
        # ép xuống dòng trước các mã số, mục lớn, số tiền
        s = re.sub(r"(?<!^)\s(?=\d{3}(?:\.\d+)?\b)", "\n", s)
        s = re.sub(r"(?<!^)\s(?=(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\.?\b)", "\n", s)
        s = re.sub(r"(?<!^)\s(?=\d{1,3}(?:[.,]\d{3}){2,}\b)", "\n", s)
        out_lines.extend([p.strip() for p in s.split("\n") if p.strip()])
    return "\n".join(out_lines)

def ocr_image_to_text_tsv(img_bgr, ocr_lang: str, ocr_cfg: str) -> str:
    try:
        cfg = (ocr_cfg or "").strip()
        cfg = re.sub(r"--psm\s+\d+", "", cfg)
        cfg = (cfg + " --psm 4 preserve_interword_spaces=1").strip()
        tsv = pytesseract.image_to_data(img_bgr, lang=ocr_lang, config=cfg, output_type=TessOutput.DICT)
        txt = reflow_lines_from_tsv_dict(tsv)
        return clean_txt_chars(txt)
    except Exception as e:
        print(f"⚠️ Lỗi OCR TSV: {e}")
        return ""

# =========================
# ⚙️ ĐỌC ẢNH & PDF SCAN
# =========================
def pdf_to_texts(pdf_path: str, dpi: int = 400,
                 ocr_lang: str = OCR_LANG_DEFAULT, ocr_cfg: str = OCR_CFG_DEFAULT,
                 start_page: Optional[int] = None, end_page: Optional[int] = None) -> List[Tuple[int, str]]:
    texts = []
    try:
        pages = convert_from_path(pdf_path, dpi=dpi)
    except Exception as e:
        print(f"⚠️ Lỗi mở PDF {pdf_path}: {e}")
        return texts

    total = len(pages)
    s = start_page or 1
    e = end_page or total
    s = max(1, s); e = min(total, e)
    for i in range(s, e+1):
        page = pages[i-1]
        img = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
        txt = ocr_image_to_text_tsv(img, ocr_lang, ocr_cfg)
        texts.append((i, txt))
    return texts

# =========================
# ⚙️ ĐỌC FILE TEXT/DOCX/EXCEL/CSV
# =========================
def read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx_file(path: str) -> str:
    if docx is None:
        raise ImportError("⚠️ Cần cài python-docx để đọc DOCX.")
    d = docx.Document(path)
    # Lấy paragraph + table (nếu có) theo dạng mô tả
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Bảng: nối theo hàng
    for tbl in d.tables:
        parts.append("========== TABLE ==========")
        for r in tbl.rows:
            cells = [clean_txt_chars(c.text) for c in r.cells]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def read_excel_or_csv(path: str, mode: str = "summary") -> str:
    ext = Path(path).suffix.lower()
    texts = []
    try:
        if ext == ".csv":
            df = pd.read_csv(path, dtype=str)
            dfs = {"CSV": df}
        else:
            xls = pd.ExcelFile(path, engine="openpyxl")
            dfs = {sheet: pd.read_excel(xls, sheet_name=sheet, dtype=str)
                   for sheet in xls.sheet_names}

        for sheet_name, df in dfs.items():
            if df.empty:
                continue
            df = df.fillna("").astype(str)
            if mode == "raw":
                headers = list(df.columns)
                header_line = " | ".join(headers)
                rows = [" | ".join(row.values) for _, row in df.iterrows()]
                sheet_text = f"\n\n========== SHEET: {sheet_name.upper()} ==========\n{header_line}\n" + "\n".join(rows)
            else:
                df = df.loc[:, ~df.columns.str.contains("^Unnamed")]
                headers = list(df.columns)
                sheet_text = [f"\n========== SHEET: {sheet_name.upper()} =========="]
                sheet_text.append(f"Các cột gồm: {', '.join(headers)}.")
                for _, row in df.iterrows():
                    if not any(v.strip() for v in row.values):
                        continue
                    pairs = [f"{h}: {v}" for h, v in row.items() if v.strip()]
                    sheet_text.append("; ".join(pairs))
                sheet_text = "\n".join(sheet_text)
            texts.append(sheet_text.strip())

        return "\n\n".join(texts).strip()
    except Exception as e:
        print(f"⚠️ Lỗi đọc {path}: {e}")
        return ""

# =========================
# ⚙️ POSTPROCESS: Anchors & chuẩn hoá
# =========================
RE_CHAPTER = re.compile(r"^\s*(CHAPTER|CHAP\.?)\s+([IVXLC]+)\b(.*)$", re.IGNORECASE)
RE_SECTION = re.compile(r"^\s*(SECTION|SEC\.?)\s+([A-Z0-9]+)\b(.*)$", re.IGNORECASE)
RE_PART    = re.compile(r"^\s*(PART|PT\.?)\s+([A-Z0-9]+)\b(.*)$", re.IGNORECASE)

def add_heading_anchors(text: str) -> Tuple[str, List[str]]:
    """
    Chuyển các heading sang Markdown anchors:
      CHAPTER I ...  -> # CHAPTER I ...
      SECTION A ...  -> ## SECTION A ...
      PART 1 ...     -> ### PART 1 ...
    Trả về (text_mới, danh_sách_anchors)
    """
    anchors: List[str] = []
    out_lines = []
    for line in (text or "").splitlines():
        raw = line.strip()
        if not raw:
            out_lines.append("")
            continue

        m = RE_CHAPTER.match(raw)
        if m:
            head = f"CHAPTER {m.group(2)}{m.group(3)}".strip()
            anchors.append(head)
            out_lines.append(f"# {head}")
            continue

        m = RE_SECTION.match(raw)
        if m:
            head = f"SECTION {m.group(2)}{m.group(3)}".strip()
            anchors.append(head)
            out_lines.append(f"## {head}")
            continue

        m = RE_PART.match(raw)
        if m:
            head = f"PART {m.group(2)}{m.group(3)}".strip()
            anchors.append(head)
            out_lines.append(f"### {head}")
            continue

        # Chuẩn hoá bullet/dash đơn giản
        line2 = re.sub(r"^\s*[-•·]\s*", "- ", raw)
        out_lines.append(line2)

    txt2 = "\n".join(out_lines)
    # Gom bớt dòng trống, loại bội khoảng trắng
    txt2 = re.sub(r"\n{3,}", "\n\n", txt2).strip()
    return txt2, sorted(list(dict.fromkeys(anchors)))  # giữ thứ tự xuất hiện

def apply_rules_and_enrich_meta(
    file_path: str, base: str, text: str, meta: dict, rules: dict
) -> Tuple[str, dict]:
    """
    - Áp YAML rules để điền company/doc_title/doc_type/...
    - Chèn anchors vào text, đưa danh sách anchors vào meta.
    """
    # 1) Anchors
    text_clean = clean_txt_chars(text)
    text_anchored, anchors = add_heading_anchors(text_clean)

    # 2) Rules → extended meta
    matched = match_rules(file_path, text_anchored, rules)
    if matched:
        meta.update({
            "company":        matched.get("company"),
            "doc_title":      matched.get("doc_title"),
            "doc_type":       matched.get("doc_type"),
            "department":     matched.get("department"),
            "lob":            matched.get("lob"),
            "appendix_id":    matched.get("appendix_id"),
            "effective_date": matched.get("effective_date"),
            "version":        matched.get("version"),
            "confidentiality":matched.get("confidentiality")
        })

    # 3) Anchors → meta
    if anchors:
        meta["anchors"] = anchors

    return text_anchored, meta

# =========================
# ⚙️ GHI OUTPUT (tôn trọng APPEND_MODE)
# =========================
def save_output_text_and_meta(text: str, meta: dict, out_txt: str, out_meta: str):
    ensure_dir(os.path.dirname(out_txt))

    # Nếu APPEND_MODE và file đã tồn tại, bỏ qua nếu sha1 không đổi
    if APPEND_MODE and os.path.exists(out_txt) and os.path.exists(out_meta):
        try:
            with open(out_txt, "r", encoding="utf-8") as f:
                old = f.read()
            old_sha = _sha1_text(old)
            new_sha = _sha1_text(text)
            if old_sha == new_sha:
                print(f"⏭️ Bỏ qua (không đổi): {os.path.basename(out_txt)}")
                return
        except Exception:
            pass

    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(text)
    with open(out_meta, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    print(f"📝 Saved: {os.path.basename(out_txt)}, {os.path.basename(out_meta)}")

# =========================
# ⚙️ XỬ LÝ FILE CHÍNH
# =========================
def process_file(file_path: str, input_root: str, output_root: str,
                 ocr_lang: str, ocr_cfg: str, rules: dict,
                 dpi: int = 400,
                 start_page: Optional[int] = None, end_page: Optional[int] = None,
                 excel_mode: str = "summary"):

    rel_path = os.path.relpath(file_path, input_root)
    base = Path(file_path).stem
    out_dir = os.path.join(output_root, os.path.dirname(rel_path))
    ensure_dir(out_dir)

    ext = Path(file_path).suffix.lower()
    text_outputs: List[Tuple[int, str]] = []

    if ext in [".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"]:
        img = cv2.cvtColor(np.array(Image.open(file_path).convert("RGB")), cv2.COLOR_RGB2BGR)
        txt = ocr_image_to_text_tsv(img, ocr_lang, ocr_cfg)
        text_outputs = [(1, txt)]

    elif ext == ".pdf":
        print(f"📄 {os.path.basename(file_path)} → OCR (PDF)")
        t0 = time.perf_counter()
        text_outputs = pdf_to_texts(file_path, dpi=dpi, ocr_lang=ocr_lang, ocr_cfg=ocr_cfg,
                                    start_page=start_page, end_page=end_page)
        t1 = time.perf_counter()
        print(f"🕓 OCR xong {len(text_outputs)} trang ({t1 - t0:.1f}s)")

    elif ext in [".doc", ".docx"]:
        try:
            txt = read_docx_file(file_path)
            text_outputs = [(1, txt)]
        except Exception as e:
            print(f"⚠️ Lỗi DOCX: {e}")

    elif ext in [".xls", ".xlsx", ".csv"]:
        txt = read_excel_or_csv(file_path, mode=excel_mode)
        text_outputs = [(1, txt)]

    elif ext == ".txt":
        txt = read_text_file(file_path)
        text_outputs = [(1, txt)]

    else:
        print(f"⚠️ Bỏ qua (định dạng không hỗ trợ): {file_path}")
        return

    # Ghi kết quả
    src_sha1 = _sha1_file(file_path)
    if len(text_outputs) > 1:
        combined_text = "\n\n".join(t for _, t in text_outputs if t.strip())
        page_range = f"{start_page or 1}-{end_page or (start_page or len(text_outputs))}"
        out_txt = os.path.join(out_dir, f"{base}_page{page_range}_text.txt")
        out_meta = os.path.join(out_dir, f"{base}_page{page_range}_meta.json")
        meta = {
            "file": base,
            "page_range": page_range,
            "page_count": len(text_outputs),
            "source_path": os.path.abspath(file_path),
            "source_doc_sha1": src_sha1,
            "language": detect_language(combined_text),
            "ocr_lang": ocr_lang,
            "ocr_cfg": ocr_cfg,
            "text_sha1": _sha1_text(combined_text),
        }
        combined_text, meta = apply_rules_and_enrich_meta(file_path, base, combined_text, meta, rules)
        save_output_text_and_meta(combined_text, meta, out_txt, out_meta)
    else:
        for page_no, txt in text_outputs:
            out_txt = os.path.join(out_dir, f"{base}_page{page_no}_text.txt")
            out_meta = os.path.join(out_dir, f"{base}_page{page_no}_meta.json")
            meta = {
                "file": base,
                "page": page_no,
                "source_path": os.path.abspath(file_path),
                "source_doc_sha1": src_sha1,
                "language": detect_language(txt),
                "ocr_lang": ocr_lang,
                "ocr_cfg": ocr_cfg,
                "text_sha1": _sha1_text(txt),
            }
            txt2, meta2 = apply_rules_and_enrich_meta(file_path, base, txt, meta, rules)
            save_output_text_and_meta(txt2, meta2, out_txt, out_meta)

# =========================
# ⚙️ MAIN ENTRYPOINT
# =========================
def main():
    parser = argparse.ArgumentParser("A1 — OCR đa định dạng (Text only, TSV reflow + YAML rules + anchors)")
    parser.add_argument("--input", type=str, default=INPUT_DIR_DEFAULT, help="Thư mục input")
    parser.add_argument("--out", type=str, default=OUTPUT_DIR_DEFAULT, help="Thư mục output")
    parser.add_argument("--rules", type=str, default=RULES_PATH_DEFAULT, help="Đường dẫn YAML rules")
    parser.add_argument("--ocr-lang", type=str, default=OCR_LANG_DEFAULT)
    parser.add_argument("--ocr-cfg", type=str, default=OCR_CFG_DEFAULT)
    parser.add_argument("--dpi", type=int, default=400)
    parser.add_argument("--clean", choices=["y","a","n","ask"], default="ask")
    parser.add_argument("--start", type=int, default=None, help="Trang bắt đầu (PDF)")
    parser.add_argument("--end", type=int, default=None, help="Trang kết thúc (PDF)")
    parser.add_argument("--excel-mode", choices=["raw", "summary"], default="summary",
                        help="Cách đọc Excel: raw=giữ nguyên, summary=làm sạch để vector store")

    args = parser.parse_args()
    START_PAGE, END_PAGE = args.start, args.end

    # Xử lý thư mục output
    if os.path.exists(args.out):
        choice = args.clean
        if choice == "ask":
            try:
                choice = input(f"⚠️ Output '{args.out}' đã tồn tại. y=xoá, a=append, n=bỏ qua: ").strip().lower()
            except Exception:
                choice = "a"  # nếu chạy không có stdin
        if choice == "y":
            shutil.rmtree(args.out, ignore_errors=True)
            print(f"🗑️ Đã xoá {args.out}")
        elif choice == "a":
            global APPEND_MODE
            APPEND_MODE = True
            print(f"➕ Giữ {args.out}, chỉ ghi file mới/khác nội dung.")
        elif choice == "n":
            print("⏭️ Bỏ qua toàn bộ."); return
        else:
            print("❌ Lựa chọn không hợp lệ."); return

    ensure_dir(args.out)
    files = glob.glob(os.path.join(args.input, "**", "*.*"), recursive=True)
    if not files:
        print("⚠️ Không tìm thấy file nào trong input."); return

    # Nạp YAML rules
    rules = load_rules(args.rules)
    if rules:
        print(f"🧩 Đã nạp rules: {args.rules}")
    else:
        print(f"ℹ️ Không tìm thấy/không đọc được rules: {args.rules} (vẫn chạy bình thường)")

    print(f"📂 Input : {args.input}")
    print(f"📦 Output: {args.out}")
    print(f"🧮 Tổng số file: {len(files)}")
    print(f"🧭 Giới hạn trang PDF: {START_PAGE or 1} → {END_PAGE or 'tất cả'}")
    has_excel = any(f.lower().endswith((".xls", ".xlsx", ".csv")) for f in files)
    if has_excel:
        print(f"📊 Chế độ đọc Excel: {args.excel_mode}")

    for f in files:
        process_file(f, args.input, args.out,
                     args.ocr_lang, args.ocr_cfg, rules,
                     dpi=args.dpi, start_page=START_PAGE, end_page=END_PAGE,
                     excel_mode=args.excel_mode)

    print("\n✅ Hoàn tất OCR. Kiểm tra *_text.txt và *_meta.json trong thư mục output.")

if __name__ == "__main__":
    main()
